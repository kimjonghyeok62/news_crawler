# Combined_News_Scraper_GUI_v5.0.1.py
# 6대 지방지 + 16개 기타 지방지 + Google 뉴스 통합 스크래퍼
#
# [v5.0.1 변경 사항]
# 1. (요청) [Step 1]에 '오늘 날짜' 목록 가져오기 버튼 추가 (날짜 동적 표시)
# 2. (요청) [Step 1]에 '엑셀파일로 다운' 버튼 추가
# 3. (요청) [Step 1] GUI는 (제목, 날짜, 키워드)로 v5.0 형식 유지
# 4. (요청) [Step 1] '엑셀 다운' 서식은 (연번, 지역, 시간, 기관 등) v1.4 형식 적용
# 5. 이를 위해 [Step 1]의 'scrape_press_releases' 함수를 v1.4 기준으로 교체 (파싱 정보 확장)
# 6. 'extract_keyword' (GUI용)와 'extract_institution' (엑셀용) 함수 동시 사용

# =============================================================================
# 1. 모든 Import 문 통합
# =============================================================================
import time
import re
import pandas as pd
import traceback
from datetime import datetime, timedelta
from urllib.parse import quote, urlencode, urljoin, urlparse, parse_qs
from bs4 import BeautifulSoup
import os
import sys
import subprocess

# --- [v5.0] Google 뉴스용 feedparser 추가 ---
import feedparser

# --- [v5.0] 보도자료 스크레이퍼용 모듈 ---
import requests
import ssl
from requests.adapters import HTTPAdapter

# --- [v5.0] Selenium 및 WebDriverManager 모듈 ---
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, WebDriverException
from webdriver_manager.chrome import ChromeDriverManager

# --- [v5.0] 병렬 처리 및 GUI 모듈 ---
from concurrent.futures import ThreadPoolExecutor, as_completed
from functools import lru_cache
import tkinter as tk
# ✅ [v5.0.1 수정] filedialog (엑셀 저장), Toplevel (팝업) 추가
from tkinter import ttk, messagebox, Toplevel, filedialog
import webbrowser
import threading

# ✅ [v5.0.1 신규] 엑셀 저장용 모듈
import xlsxwriter
# ✅ Word 문서 생성을 위한 모듈
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE

# =============================================================================
# 2. 공통 헬퍼 함수 (보도자료 스크레이퍼)
# =============================================================================

class LegacyTLSAdapter(HTTPAdapter):
    def init_poolmanager(self, *args, **kwargs):
        context = ssl.create_default_context(ssl.Purpose.SERVER_AUTH)
        try:
            context.set_ciphers('DEFAULT@SECLEVEL=1')
        except ssl.SSLError:
            print("경고: set_ciphers('DEFAULT@SECLEVEL=1') 설정에 실패했습니다. 기본값으로 진행합니다.")
        
        kwargs['ssl_context'] = context
        return super(LegacyTLSAdapter, self).init_poolmanager(*args, **kwargs)

# ✅ [v5.0.1] v5.0의 'extract_keyword' (GUI용) 함수 유지
def extract_keyword(title):
    try:
        # ✅ [v4.0] 특정 학교 이름이 제목에 포함된 경우, 해당 이름을 최우선 반환
        specific_schools = ["한사랑학교", "성광학교", "광주새롬학교", "동현학교", "인덕학교"]
        for school in specific_schools:
            if school in title:
                return school # (1순위) 특정 학교 이름 반환

        # (이하 기존 로직 - 2순위)
        if "광주하남교육지원청" in title:
            return "광주하남교육"
        
        # (이하 기존 로직 - 3순위)
        match = re.search(r'\b([\w]+(초|중|고|유|병설유|초병설유|병설유치원|초등학교|중학교|고등학교|유치원))\b', title)
        
        if match:
            keyword = match.group(1)
            return keyword
        
        # (이하 기존 로직 - 4순위)
        first_part = title.split(',')[0].strip()
        first_word = first_part.split(' ')[0].strip()
        
        if len(first_word) > 10:
            return first_word[:10]
        
        return first_word

    except Exception as e:
        print(f"키워드 추출 오류: {e} (제목: {title})")
        return title[:5]

# ✅ [v5.0.1 신규] 엑셀 저장용 '등록기관' 추출 함수 (from 뷰어 v1.4)
def extract_institution(title):
    """ (v1.4 로직) """
    try:
        specific_schools = ["한사랑학교", "성광학교", "광주새롬학교", "동현학교", "인덕학교"]
        for school in specific_schools:
            if school in title:
                return school 
        if "광주하남교육지원청" in title:
            return "광주하남교육"
        match = re.search(r'\b([\w]+(초|중|고|유|병설유|초병설유|병설유치원|초등학교|중학교|고등학교|유치원))\b', title)
        if match:
            return match.group(1)
        first_part = title.split(',')[0].strip()
        first_word = first_part.split(' ')[0].strip()
        return first_word[:10] if len(first_word) > 10 else first_word
    except Exception as e:
        print(f"등록기관 추출 오류: {e} (제목: {title})")
        return title[:5]


# ✅ [v5.0.1 수정] v1.4의 엑셀 저장용 정보까지 파싱하도록 함수 교체
def scrape_press_releases(base_url):
    """
    [v5.0.1 수정]
    지정된 base_url의 1페이지에서 보도자료를 스크래핑합니다.
    - GUI용 'keyword_gui' (from extract_keyword)
    - 엑셀용 'institution', 'region' 등 (from extract_institution)
    - [요청 1] 등록시간 파싱 제거.
    """
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
    }
    session = requests.Session()
    session.mount("https://", LegacyTLSAdapter())
    session.mount("http://", LegacyTLSAdapter())
    results_list = []
    total_articles_found = 0

    try:
        # 1페이지만 수집
        page_url = f"{base_url}&pageIndex=1"
        print(f"[보도자료] 1페이지 스크래핑... (URL: {page_url})")

        response = session.get(page_url, headers=headers, timeout=10)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.text, 'html.parser')
        articles = soup.select('a.selectNttInfo')

        if not articles:
            print(f"[보도자료] 1페이지에서 게시글을 찾지 못했습니다.")
            return []

        for article in articles:
            original_title = article.get('title', '제목 없음').strip()
            
            # --- 날짜/시간 파싱 (v1.4 기준) ---
            date_tag = article.find('span', class_='date')
            date_str_raw = date_tag.get_text(strip=True) if date_tag else '날짜 없음'
            date = "날짜없음"
            time = "시간없음" # (v1.4) 스레드 실행 시간으로 덮어쓸 예정
            
            date_match = re.search(r"(\d{4}\.\d{2}\.\d{2})", date_str_raw)
            if date_match:
                date = date_match.group(0)
            else:
                print(f"⚠️ [보도자료] 날짜 형식 불일치로 건너뜀: {date_str_raw}")
                continue

            # --- 링크 파싱 (v1.4 기준) ---
            nttSn = article.get('data-param') 
            link = "링크 추출 실패" 
            if nttSn and nttSn.isdigit():
                link = f"https://www.goegh.kr/goegh/na/ntt/selectNttInfo.do?mi=8686&bbsId=5041&nttSn={nttSn}"
            else:
                print(f"⚠️ [보도자료] 링크(nttSn) 추출 실패. (data-param 값: {nttSn})")

            # --- 제목/지역/기관 파싱 (v1.4 기준) ---
            clean_title = original_title 
            title_match = re.search(r'^보도자료\((.*)\)$', original_title)
            if title_match:
                clean_title = title_match.group(1).strip() 
            
            region = "교육지원청"
            if clean_title.startswith("광주 "): region = "광주"
            elif clean_title.startswith("하남 "): region = "하남"
            elif "광주하남" in clean_title: region = "교육지원청"

            # ✅ 엑셀 저장용 (v1.4)
            institution = extract_institution(clean_title)
            # ✅ GUI 표시용 (v5.0)
            keyword_gui = extract_keyword(clean_title) 
            
            results_list.append({
                "priority": "",
                "region": region,
                "date": date,
                "time": time,
                "title": clean_title,
                "institution": institution, # 엑셀용
                "keyword_gui": keyword_gui, # GUI용
                "link": link, 
                "notes": ""
            })
            total_articles_found += 1
        
        print(f"총 {total_articles_found}개의 게시글을 찾았습니다. (1 페이지)")
        return results_list

    except requests.exceptions.RequestException as e:
        print(f"HTTP 요청 중 오류가 발생했습니다: {e}")
    except Exception as e:
        print(f"알 수 없는 오류가 발생했습니다: {e}")
        traceback.print_exc()
    return []

# =============================================================================
# 3. [v5.0] Selenium 드라이버 설정 (6대, 16대 공통 사용)
# (이하 v5.0 원본과 동일 ... setup_driver, setup_driver_compatible)
# =============================================================================

USER_AGENT = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"

@lru_cache(maxsize=1)
def get_driver_path():
    print("WebDriverManager: 크롬 드라이버 설치 또는 업데이트 확인 중...")
    try:
        path = ChromeDriverManager().install()
        print(f"WebDriverManager: 드라이버 경로 확인: {path}")
        return path
    except Exception as e:
        print(f"WebDriverManager: 드라이버 설치 실패: {e}")
        print("PATH에 chromedriver가 수동으로 설치되어 있는지 확인하세요.")
        return "chromedriver"

def setup_driver(headless=True):
    """
    최적화된 옵션으로 크롬 웹 드라이버를 설정합니다. (대부분의 사이트용)
    (from '기타지방지 크롤러' - prefs 포함 버전)
    """
    options = Options()
    if headless:
        options.add_argument("--headless=new")
    
    options.add_argument(f"--user-agent={USER_AGENT}")
    options.add_argument("--disable-gpu")
    options.add_argument("--no-sandbox")
    options.add_argument("--window-size=1920,1080")
    options.add_argument("--log-level=3")
    options.add_experimental_option('excludeSwitches', ['enable-logging'])
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--ignore-certificate-errors")
    options.add_argument("--ignore-ssl-errors")
    options.add_argument("--disable-extensions")
    options.add_argument("--disable-popup-blocking")
    options.add_argument("--disable-infobars")
    options.add_argument("--disable-component-extensions-with-background-pages")
    options.add_argument("--disable-default-apps")
    options.add_argument("--disable-sync")

    # [v5.0] '기타지방지'의 리소스 최적화(prefs) 옵션 사용
    prefs = {
        "profile.managed_default_content_settings.images": 2,
        "profile.managed_default_content_settings.stylesheets": 2,
        "profile.managed_default_content_settings.cookies": 2,
        "profile.managed_default_content_settings.javascript": 1,
        "profile.managed_default_content_settings.plugins": 2,
        "profile.managed_default_content_settings.popups": 2,
        "profile.managed_default_content_settings.geolocation": 2,
        "profile.managed_default_content_settings.notifications": 2,
        "profile.managed_default_content_settings.media_stream": 2,
    }
    options.add_experimental_option("prefs", prefs)
    
    try:
        service = Service(get_driver_path()) 
        return webdriver.Chrome(service=service, options=options)
    except Exception as e:
        print(f"드라이버 시작 오류: {e}")
        print("드라이버 버전이 크롬 브라우저 버전과 맞는지 확인하세요.")
        return None

def setup_driver_compatible(headless=True):
    """
    호환성을 위해 리소스 차단(prefs) 옵션을 제거한 드라이버입니다.
    (from '기타지방지 크롤러')
    """
    options = Options()
    if headless:
        options.add_argument("--headless=new")
    
    options.add_argument(f"--user-agent={USER_AGENT}")
    options.add_argument("--disable-gpu")
    options.add_argument("--no-sandbox")
    options.add_argument("--window-size=1920,1080")
    options.add_argument("--log-level=3")
    options.add_experimental_option('excludeSwitches', ['enable-logging'])
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--ignore-certificate-errors")
    options.add_argument("--ignore-ssl-errors")
    options.add_argument("--disable-extensions")
    options.add_argument("--disable-popup-blocking")
    options.add_argument("--disable-infobars")
    options.add_argument("--disable-component-extensions-with-background-pages")
    options.add_argument("--disable-default-apps")
    options.add_argument("--disable-sync")
    
    try:
        service = Service(get_driver_path()) 
        return webdriver.Chrome(service=service, options=options)
    except Exception as e:
        print(f"드라이버 시작 오류: {e}")
        print("드라이버 버전이 크롬 브라우저 버전과 맞는지 확인하세요.")
        return None

# =============================================================================
# 4. [v5.0] Fetcher 함수 그룹 1 (주요 6대 지방지)
# (이하 v5.0 원본과 동일 ... fetch_kiho_multi ~ fetch_joongbu_multi)
# =============================================================================

# --- 기호일보 스크레이퍼 (fetch_kiho_multi) ---
def fetch_kiho_multi(keywords, date_limit, days_limit):
    
    ARTICLE_LIST_SELECTOR = "li.altlist-text-item"
    TITLE_SELECTOR = "h2.altlist-subject a"
    DATE_SELECTOR = "div.altlist-info div.altlist-info-item"
    IS_DATE_IN_LIST = True 
    DATE_INDEX = 2 
    
    source_name = "기호일보"
    
    driver = setup_driver(headless=True)
    if not driver: return []
    results = []
    
    try:
        for keyword in keywords:
            url = f"https://www.kihoilbo.co.kr/news/articleList.html?sc_word={quote(keyword)}&sc_order_by=E"
            print(f"🌐 [{source_name}] '{keyword}' 검색 중... (URL: {url})")
            
            driver.get(url)
            
            try:
                WebDriverWait(driver, 10).until( 
                    EC.presence_of_element_located((By.CSS_SELECTOR, ARTICLE_LIST_SELECTOR))
                )
                
                soup = BeautifulSoup(driver.page_source, "html.parser")
                articles = soup.select(ARTICLE_LIST_SELECTOR)
                found_count = 0
                
                for item in articles:
                    title = "" 
                    try:
                        title_tag = item.select_one(TITLE_SELECTOR)
                        if not title_tag: continue 
                        title = title_tag.get_text(strip=True)
                        if not title: continue 
                        link = title_tag["href"]
                        if not link.startswith("http"):
                            link = urljoin("https://www.kihoilbo.co.kr", link) 
                        
                        date_str_raw = "" 
                        if IS_DATE_IN_LIST:
                            info_items = item.select(DATE_SELECTOR)
                            if len(info_items) <= DATE_INDEX:
                                continue
                            date_str_raw = info_items[DATE_INDEX].get_text(strip=True) 
                        else:
                            date_tag = item.select_one(DATE_SELECTOR)
                            if not date_tag: continue
                            date_str_raw = date_tag.get_text(strip=True)
                        
                        if not date_str_raw: continue 

                        date_str = date_str_raw.split(" ")[0] 
                        
                        if re.match(r"^\d{4}\.\d{2}\.\d{2}$", date_str):
                            date_obj = datetime.strptime(date_str, "%Y.%m.%d")
                        elif re.match(r"^\d{4}-\d{2}-\d{2}$", date_str):
                            date_obj = datetime.strptime(date_str, "%Y-%m-%d")
                        else:
                            continue
                        
                        if date_obj >= date_limit:
                            results.append({"보도일": date_obj.strftime("%Y-%m-%d"), "보도매체": source_name, "보도제목": title, "링크": link, "검색어": keyword})
                            found_count += 1
                    except Exception as e:
                        print(f"[{source_name}] 기사 1개 파싱 오류: {e} (제목: {title})")
                        continue
                
                if found_count == 0:
                    results.append({
                        "보도일": "없음", "보도매체": source_name,
                        "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                        "링크": url, "검색어": keyword
                    })

            except TimeoutException as wait_e: 
                print(f"[{source_name}] '{keyword}' 검색 결과 없음 (타임아웃).")
                results.append({
                    "보도일": "없음", "보도매체": source_name,
                    "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                    "링크": url, "검색어": keyword
                })

    except Exception as e:
        print(f"[{source_name}] 페이지 로드 또는 드라이버 오류: {e}")
    finally:
        if driver:
            driver.quit()
    return results


# --- 인천일보 스크레이퍼 (fetch_incheonilbo_multi) ---
def fetch_incheonilbo_multi(keywords, date_limit, days_limit):
    source_name = "인천일보"
    base_url = "https://www.incheonilbo.com"
    
    driver = setup_driver(headless=True)
    if not driver: return []
    results = []
    
    try:
        for keyword in keywords:
            url = f"https://www.incheonilbo.com/news/articleList.html?sc_word={quote(keyword)}&sc_order_by=E"
            print(f"🌐 [{source_name}] '{keyword}' 검색 중... (URL: {url})")
            
            driver.get(url)
            
            try:
                WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.CSS_SELECTOR, "ul.type1 > li"))
                )
                
                soup = BeautifulSoup(driver.page_source, "html.parser")
                articles = soup.select("ul.type1 > li")
                found_count = 0

                for article in articles:
                    title = ""
                    try:
                        title_tag = article.select_one("h2.titles")
                        link_tag = article.select_one("a")
                        if not title_tag or not link_tag: continue
                        title = title_tag.text.strip()
                        if not title: continue 
                        link = urljoin(base_url, link_tag["href"])
                        date_str_tag = article.select_one("em.info.dated")
                        pub_date = None

                        if date_str_tag:
                            date_text = date_str_tag.text.strip()
                            if not date_text: continue 
                            try:
                                pub_date = datetime.strptime(date_text.split(" ")[0], "%Y.%m.%d")
                            except ValueError:
                                if "분 전" in date_text:
                                    minutes_ago = int(re.search(r'(\d+)', date_text).group(1))
                                    pub_date = datetime.now() - timedelta(minutes=minutes_ago)
                                elif "시간 전" in date_text:
                                    hours_ago = int(re.search(r'(\d+)', date_text).group(1))
                                    pub_date = datetime.now() - timedelta(hours=hours_ago)
                                else:
                                    continue
                        else:
                            continue

                        if pub_date and pub_date >= date_limit:
                            results.append({"보도일": pub_date.strftime("%Y-%m-%d"), "보도매체": source_name, "보도제목": title, "링크": link, "검색어": keyword})
                            found_count += 1
                    except Exception as e:
                        print(f"[{source_name}] 기사 파싱 오류: {e} (제목: {title})")
                        continue
                
                if found_count == 0:
                    results.append({
                        "보도일": "없음", "보도매체": source_name,
                        "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                        "링크": url, "검색어": keyword
                    })

            except TimeoutException as wait_e: 
                print(f"[{source_name}] '{keyword}' 검색 결과 없음 (타임아웃).")
                results.append({
                    "보도일": "없음", "보도매체": source_name,
                    "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                    "링크": url, "검색어": keyword
                })
            
    except Exception as e:
        print(f"[{source_name}] 페이지 로드 오류: {e}")
    finally:
        if driver:
            driver.quit()
    return results

# --- 경기일보 스크레이퍼 (fetch_kyeonggi_multi) ---
def fetch_kyeonggi_multi(keywords, date_limit, days_limit):
    source_name = "경기일보"
    base_url = "https://www.kyeonggi.com"
    
    driver = setup_driver(headless=True)
    if not driver: return []
    results = []
    
    try:
        for keyword in keywords:
            url = f"https://www.kyeonggi.com/search?searchText={quote(keyword)}"
            print(f"🌐 [{source_name}] '{keyword}' 검색 중... (URL: {url})")

            driver.get(url)
            
            try:
                WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.CSS_SELECTOR, "div.article_list div.media"))
                )
                
                soup = BeautifulSoup(driver.page_source, "html.parser")
                articles = soup.select("div.article_list div.media")
                found_count = 0
                
                for article in articles: 
                    title = ""
                    date_str = ""
                    try:
                        title_tag = article.select_one("h3 a")
                        if not title_tag: continue
                        title = title_tag.get_text(strip=True)
                        if not title: continue 
                        link = urljoin(base_url, title_tag.get("href", "")) 
                        date_tag = article.select_one("span.byline") 
                        if not date_tag: continue
                        
                        date_text_raw = date_tag.get_text() 
                        if not date_text_raw: continue

                        date_match = re.search(r"(\d{4}\.\d{2}\.\d{2})", date_text_raw)
                        if date_match:
                            date_str = date_match.group(1).replace(".", "-")
                        else:
                            date_elements = date_tag.select("span")
                            if len(date_elements) >= 3:
                                date_str_raw = date_elements[2].get_text(strip=True)
                            elif len(date_elements) >= 2:
                                date_str_raw = date_elements[1].get_text(strip=True)
                            else:
                                continue
                            date_str = date_str_raw.split(" ")[0].replace(".", "-")
                        
                        if not re.match(r"^\d{4}-\d{2}-\d{2}$", date_str):
                            continue
                        
                        pub_date = datetime.strptime(date_str, "%Y-%m-%d")
                        if pub_date >= date_limit:
                            results.append({"보도일": pub_date.strftime("%Y-%m-%d"), "보도매체": source_name, "보도제목": title, "링크": link, "검색어": keyword})
                            found_count += 1
                    except Exception as e:
                        print(f"[{source_name}] 기사 파싱 오류: {e} (날짜: {date_str}, 제목: {title})")
                        continue
                
                if found_count == 0:
                    results.append({
                        "보도일": "없음", "보도매체": source_name,
                        "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                        "링크": url, "검색어": keyword
                    })
                    
            except TimeoutException as wait_e: 
                print(f"[{source_name}] '{keyword}' 검색 결과 없음 (타임아웃).")
                results.append({
                    "보도일": "없음", "보도매체": source_name,
                    "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                    "링크": url, "검색어": keyword
                })
                
    except Exception as e:
        print(f"[{source_name}] 페이지 로드 오류: {e}")
    finally:
        if driver:
            driver.quit()
    return results

# --- 경인일보 스크레이퍼 (fetch_kyeongin_multi) ---
def fetch_kyeongin_multi(keywords, date_limit, days_limit):
    source_name = "경인일보"
    base_url = "https://www.kyeongin.com"
    
    driver = setup_driver(headless=True)
    if not driver: return []
    results = []
    
    try:
        for keyword in keywords:
            url = f"https://www.kyeongin.com/search?query={quote(keyword)}"
            print(f"🌐 [{source_name}] '{keyword}' 검색 중... (URL: {url})")

            driver.get(url)
            try:
                WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.CSS_SELECTOR, "div.search-arl-001 ul li"))
                )
                
                soup = BeautifulSoup(driver.page_source, "html.parser")
                articles = soup.select("div.search-arl-001 ul li")
                found_count = 0
                
                for item in articles:
                    title = ""
                    try:
                        a_tag = item.select_one("h4.title a")
                        if not a_tag: continue
                        title = a_tag.get_text(strip=True)
                        if not title: continue 
                        
                        href = a_tag["href"]
                        if href.startswith("//"):
                            link = "https:" + href
                        else:
                            link = urljoin(base_url, href)
                        
                        date_tag = item.select_one("span.date") 
                        if not date_tag: continue
                        date_str = date_tag.get_text(strip=True)
                        if not date_str: continue

                        pub_date = datetime.strptime(date_str, "%Y-%m-%d")
                        if pub_date >= date_limit:
                            results.append({"보도일": pub_date.strftime("%Y-%m-%d"), "보도매체": source_name, "보도제목": title, "링크": link, "검색어": keyword})
                            found_count += 1
                    except Exception as e:
                        print(f"[{source_name}] 기사 파싱 오류: {e} (제목: {title})")
                        continue
                
                if found_count == 0:
                    results.append({
                        "보도일": "없음", "보도매체": source_name,
                        "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                        "링크": url, "검색어": keyword
                    })
                    
            except TimeoutException as wait_e: 
                print(f"[{source_name}] '{keyword}' 검색 결과 없음 (타임아웃).")
                results.append({
                    "보도일": "없음", "보도매체": source_name,
                    "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                    "링크": url, "검색어": keyword
                })
                
    except Exception as e:
        print(f"[{source_name}] 페이지 로드 오류: {e}")
    finally:
        if driver:
            driver.quit()
    return results

# --- 경기신문 스크레이퍼 (fetch_kgnews_multi) ---
def fetch_kgnews_multi(keywords, date_limit, days_limit):
    source_name = "경기신문"
    base_url = "https://www.kgnews.co.kr"
    
    driver = setup_driver(headless=True)
    if not driver: return []
    results = []
    
    try:
        for keyword in keywords:
            url = f"https://www.kgnews.co.kr/news/search_result.html?search_mode=multi&s_title=1&s_sub_title=1&s_body=1&search={quote(keyword)}"
            print(f"🌐 [{source_name}] '{keyword}' 검색 중... (URL: {url})")

            driver.get(url)
            
            try:
                WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.CSS_SELECTOR, "ul.art_list_all > li"))
                )
                
                soup = BeautifulSoup(driver.page_source, "html.parser")
                articles = soup.select("ul.art_list_all > li")
                found_count = 0
                
                for article in articles:
                    title = ""
                    try:
                        title_tag = article.select_one("h2.clamp.c2")
                        link_tag = article.select_one("a")
                        info_tag = article.select_one("ul.ffd.art_info")
                        if not title_tag or not link_tag or not info_tag: continue
                        title = title_tag.text.strip()
                        if not title: continue 
                        link = urljoin(base_url, link_tag["href"]) 
                        date_text_raw = info_tag.text 
                        if not date_text_raw: continue
                        date_match = re.search(r"\d{4}\.\d{2}\.\d{2}", date_text_raw)
                        if not date_match: continue

                        pub_date = datetime.strptime(date_match.group(), "%Y.%m.%d")
                        if pub_date >= date_limit:
                            results.append({"보도일": pub_date.strftime("%Y-%m-%d"), "보도매체": source_name, "보도제목": title, "링크": link, "검색어": keyword})
                            found_count += 1
                    except Exception as e:
                        print(f"[{source_name}] 기사 파싱 오류: {e} (제목: {title})")
                        continue

                if found_count == 0:
                    results.append({
                        "보도일": "없음", "보도매체": source_name,
                        "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                        "링크": url, "검색어": keyword
                    })
                    
            except TimeoutException as wait_e: 
                print(f"[{source_name}] '{keyword}' 검색 결과 없음 (타임아웃).")
                results.append({
                    "보도일": "없음", "보도매체": source_name,
                    "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                    "링크": url, "검색어": keyword
                })
                
    except Exception as e:
        print(f"[{source_name}] 페이지 로드 오류: {e}")
    finally:
        if driver:
            driver.quit()
    return results

# --- 중부일보 스크레이퍼 (fetch_joongbu_multi) ---
def fetch_joongbu_multi(keywords, date_limit, days_limit):
    source_name = "중부일보"
    base_url = "https://www.joongboo.com"
    
    driver = setup_driver(headless=True)
    if not driver: return []
    results = []

    try:
        for keyword in keywords:
            url = f"https://www.joongboo.com/news/articleList.html?sc_area=A&sc_word={quote(keyword)}&sc_order_by=E"
            print(f"🌐 [{source_name}] '{keyword}' 검색 중... (URL: {url})")

            driver.get(url)
            
            try:
                WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.CSS_SELECTOR, "div.list-content"))
                )
                
                soup = BeautifulSoup(driver.page_source, "html.parser")
                articles = soup.select("div.list-content")
                found_count = 0
                
                for article in articles:
                    title = ""
                    try:
                        title_tag = article.select_one("h4.titles a")
                        if not title_tag: continue
                        title = title_tag.get_text(strip=True)
                        if not title: continue 
                        link = urljoin(base_url, title_tag["href"])
                        date_elements = article.select("span.byline em")
                        if len(date_elements) < 2: continue
                        date_str = date_elements[1].text.strip().split(" ")[0]
                        if not date_str: continue

                        pub_date = datetime.strptime(date_str, "%Y.%m.%d")
                        if pub_date >= date_limit:
                            results.append({"보도일": pub_date.strftime("%Y-%m-%d"), "보도매체": source_name, "보도제목": title, "링크": link, "검색어": keyword})
                            found_count += 1
                    except Exception as e:
                        print(f"[{source_name}] 기사 파싱 오류: {e} (제목: {title})")
                        continue
                
                if found_count == 0:
                    results.append({
                        "보도일": "없음", "보도매체": source_name,
                        "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                        "링크": url, "검색어": keyword
                    })
                    
            except TimeoutException as wait_e: 
                print(f"[{source_name}] '{keyword}' 검색 결과 없음 (타임아웃).")
                results.append({
                    "보도일": "없음", "보도매체": source_name,
                    "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                    "링크": url, "검색어": keyword
                })
                
    except Exception as e:
        print(f"[{source_name}] 페이지 로드 오류: {e}")
    finally:
        if driver:
            driver.quit()
    return results

# =============================================================================
# 5. [v5.0] Fetcher 함수 그룹 2 (기타 16개 지방지)
# (이하 v5.0 원본과 동일 ... fetch_kghottimes_multi ~ fetch_sudokwon_multi)
# =============================================================================

# fetch_kghottimes.py (최적화)
def fetch_kghottimes_multi(keywords, date_limit, days_limit): 
    results = []
    driver = setup_driver(headless=True)
    if not driver: return []
    try:
        base_url = "http://ghottimenews.com/news/search_result.html"
        
        for keyword in keywords:
            search_url = f"{base_url}?search_mode=multi&s_title=1&s_body=1&search={quote(keyword)}"
            print(f"🌐 [경기핫타임스] '{keyword}' 검색 중... (URL: {search_url})")
            
            driver.get(search_url)
            
            try:
                WebDriverWait(driver, 5).until(
                    EC.presence_of_all_elements_located((By.CSS_SELECTOR, "ul.art_list_all li"))
                )
                soup = BeautifulSoup(driver.page_source, "html.parser")
                found_count = 0
                for li in soup.select("ul.art_list_all li"):
                    try:
                        title_tag = li.select_one("h2")
                        a_tag = li.select_one("a[href*='article.html']")
                        date_tag = li.select_one("li.date")
                        
                        if not title_tag or not a_tag or not date_tag: continue
                        
                        title = title_tag.get_text(strip=True)
                        link = "http://ghottimenews.com/news/" + a_tag["href"].split("/")[-1]
                        
                        date_match = re.search(r"\d{4}\.\d{2}\.\d{2}", date_tag.get_text(strip=True))
                        if not date_match: continue
                        
                        pub_date = datetime.strptime(date_match.group(), "%Y.%m.%d")
                        
                        if pub_date >= date_limit:
                            results.append({"보도일": pub_date.strftime("%Y-%m-%d"), "보도매체": "경기핫타임스", "보도제목": title, "링크": link, "검색어": keyword})
                            found_count += 1
                    except Exception as e_item:
                        print(f"⚠️ [경기핫타임] 개별 파싱 예외: {e_item} (키워드: {keyword})")
                        continue
                
                if found_count == 0:
                    results.append({
                        "보도일": f"{days_limit}일 이내 없음",
                        "보도매체": "경기핫타임스",
                        "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                        "링크": search_url,
                        "검색어": keyword
                    })
                
            except TimeoutException:
                print(f"🔍 [경기핫타임스] '{keyword}' 검색 결과 없음 (타임아웃)")
                results.append({
                    "보도일": f"{days_limit}일 이내 없음",
                    "보도매체": "경기핫타임스",
                    "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                    "링크": search_url,
                    "검색어": keyword
                })

    except Exception as e_page:
        print(f"❌ [경기핫타임] 페이지 로딩 실패: {e_page}")
    finally:
        if driver:
            try: driver.quit() 
            except Exception as e_quit: print(f"⚠️ [경기핫타임스] 드라이버 종료 예외: {e_quit}")
    return results

# fetch_hanamtimes.py (최적화)
def fetch_hanamtimes_multi(keywords, date_limit, days_limit):
    results = []
    base_url = "http://www.hanamtimes.com"
    driver = setup_driver_compatible(headless=True) # 호환성 드라이버 사용
    if not driver: return []
    
    try:
        for keyword in keywords:
            encoded_keyword = quote(keyword.encode('euc-kr'))
            search_url = f"{base_url}/news/articleList.html?sc_sub_section_code=S2N11&view_type=sm&sc_word={encoded_keyword}"
            print(f"🌐 [하남타임즈] '{keyword}' 검색 중... (URL: {search_url})")

            driver.get(search_url)

            try:
                WebDriverWait(driver, 5).until(
                    EC.presence_of_all_elements_located((By.CSS_SELECTOR, "td.list-titles"))
                )
                
                soup = BeautifulSoup(driver.page_source, "html.parser")
                articles = soup.select("td.list-titles")
                found_count = 0
                seen_links = set()
                
                for title_td in articles:
                    try:
                        a_tag = title_td.select_one("a")
                        if not a_tag: continue
                        
                        title = a_tag.get_text(strip=True)
                        href = a_tag.get("href")
                        link = urljoin(search_url, href)
                        
                        if not link or link in seen_links: continue
                        seen_links.add(link)
                        
                        parent_tbody = a_tag.find_parent("tbody")
                        if not parent_tbody: continue
                        
                        date_tag = parent_tbody.select_one("td.list-times")
                        if not date_tag: continue
                        
                        date_text = date_tag.get_text(strip=True)
                        date_match = re.search(r"(\d{4}-\d{2}-\d{2})", date_text)
                        
                        if date_match:
                            pub_date = datetime.strptime(date_match.group(1), "%Y-%m-%d")
                            if pub_date >= date_limit:
                                results.append({"보도일": pub_date.strftime("%Y-%m-%d"), "보도매체": "하남타임즈", "보도제목": title, "링크": link, "검색어": keyword})
                                found_count += 1
                        
                    except Exception as e_item:
                        print(f"⚠️ [하남타임즈] 기사 파싱 오류: {e_item} (키워드: {keyword})")
                        continue

                if found_count == 0:
                    results.append({
                        "보도일": f"{days_limit}일 이내 없음",
                        "보도매체": "하남타임즈",
                        "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다. (검색어 링크 확인)",
                        "링크": search_url,
                        "검색어": keyword
                    })
                    
            except TimeoutException:
                print(f"🔍 [하남타임즈] '{keyword}' 검색 결과 없음 (타임아웃)")
                results.append({
                    "보도일": "타임아웃",
                    "보도매체": "하남타임즈",
                    "보도제목": f"웹사이트 응답 지연으로 검색 실패 (키워드: '{keyword}')",
                    "링크": search_url,
                    "검색어": keyword
                })
                
    except Exception as e_page:
        print(f"❌ [하남타임즈] 드라이버 실패: {e_page}")
    finally:
        if driver:
            try: driver.quit()
            except Exception as e_quit: print(f"⚠️ [하남타임즈] 드라이버 종료 예외: {e_quit}")
    return results

# fetch_ctnews.py (최적화)
def fetch_ctnews_multi(keywords, date_limit, days_limit):
    results = []
    base_url = "https://www.ctnews.co.kr"
    driver = setup_driver(headless=True)
    if not driver: return []
    
    try:
        for keyword in keywords:
            encoded_keyword = quote(keyword.encode("euc-kr"))
            search_url = f"{base_url}/search.html?submit=submit&search={encoded_keyword}&search_and=1&search_exec=n_b&search_section=all&news_order=1"
            print(f"🌐 [시티뉴스] '{keyword}' 검색 중... (URL: {search_url})")

            driver.get(search_url)

            try:
                WebDriverWait(driver, 5).until(
                    EC.presence_of_all_elements_located((By.CSS_SELECTOR, "td[style*='padding'] a"))
                )
                
                soup = BeautifulSoup(driver.page_source, "html.parser")
                rows = soup.select("td[style*='padding']")
                found_count = 0
                
                for row in rows:
                    try:
                        title_tag = row.select_one("a[href*='sub_read.html']")
                        date_tag = row.find_next("td", class_="data")
                        
                        if not (title_tag and date_tag): continue
                        
                        title = title_tag.get_text(strip=True)
                        href = title_tag.get("href", "")
                        full_link = f"{base_url}/{href.lstrip('/')}"
                        raw_date_text = date_tag.get_text(strip=True)
                        
                        date_match = re.search(r"\d{4}[/-]\d{2}[/-]\d{2}", raw_date_text)
                        if not date_match: continue
                        
                        pub_date = datetime.strptime(date_match.group().replace("/", "-"), "%Y-%m-%d")
                        
                        if pub_date >= date_limit:
                            results.append({"보도일": pub_date.strftime("%Y-%m-%d"), "보도매체": "시티뉴스", "보도제목": title, "링크": full_link, "검색어": keyword})
                            found_count += 1
                    except Exception as e_item:
                        print(f"⚠️ [시티뉴스] 파싱 예외: {e_item} (키워드: {keyword})")
                        continue

                if found_count == 0:
                    results.append({
                        "보도일": f"{days_limit}일 이내 없음",
                        "보도매체": "시티뉴스",
                        "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                        "링크": search_url,
                        "검색어": keyword
                    })
                    
            except TimeoutException:
                print(f"🔍 [시티뉴스] '{keyword}' 검색 결과 없음 (타임아웃)")
                results.append({
                    "보도일": f"{days_limit}일 이내 없음",
                    "보도매체": "시티뉴스",
                    "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                    "링크": search_url,
                    "검색어": keyword
                })
                
    except Exception as e_page:
        print(f"❌ [시티뉴스] 로딩 실패: {e_page}")
    finally:
        if driver:
            try: driver.quit()
            except Exception as e_quit: print(f"⚠️ [시티뉴스] 드라이버 종료 예외: {e_quit}")
    return results

# fetch_goodtimes.py (최적화 + 날짜 파싱 오류 수정)
def fetch_goodtimes_multi(keywords, date_limit, days_limit):
    results = []
    base_url = "https://www.goodtimes.or.kr"
    driver = setup_driver(headless=True)
    if not driver: return []
    
    try:
        current_year = datetime.today().year
        for keyword in keywords:
            search_url = f"{base_url}/news/articleList.html?sc_area=A&view_type=sm&sc_word={quote(keyword)}"
            print(f"🌐 [굿타임즈] '{keyword}' 검색 중... (URL: {search_url})")

            driver.get(search_url)

            try:
                WebDriverWait(driver, 5).until(
                    EC.presence_of_all_elements_located((By.CSS_SELECTOR, "li.altlist-webzine-item"))
                )
                
                soup = BeautifulSoup(driver.page_source, "html.parser")
                rows = soup.select("li.altlist-webzine-item")
                found_count = 0
                
                for row in rows:
                    try:
                        a_tag = row.select_one("h2.altlist-subject a")
                        date_tag_all = row.select("div.altlist-info-item")
                        
                        if not a_tag or not date_tag_all: continue
                        
                        title = a_tag.get_text(strip=True)
                        full_link = a_tag.get("href")
                        
                        pub_date = None
                        
                        for tag in date_tag_all:
                            date_text = tag.get_text(strip=True)
                            match = re.search(r"(\d{4})[.-](\d{2})[.-](\d{2})", date_text)
                            if match:
                                try:
                                    pub_date = datetime.strptime(match.group(0).replace(".", "-"), "%Y-%m-%d")
                                    break 
                                except Exception:
                                    continue
                        
                        if not pub_date:
                            for tag in date_tag_all:
                                date_text = tag.get_text(strip=True)
                                match = re.search(r"(\d{2})-(\d{2})", date_text) # MM-DD
                                if match:
                                    try:
                                        month, day = match.group(1), match.group(2)
                                        if 1 <= int(month) <= 12 and 1 <= int(day) <= 31:
                                            full_date_str = f"{current_year}-{month}-{day}"
                                            pub_date = datetime.strptime(full_date_str, "%Y-%m-%d")
                                            break
                                    except Exception:
                                        continue 
                        
                        if not pub_date:
                            print(f"⚠️ [굿타임즈] 날짜 파싱 실패 (제목: {title[:20]}...)")
                            continue
                            
                        if pub_date >= date_limit:
                            results.append({"보도일": pub_date.strftime("%Y-%m-%d"), "보도매체": "굿타임즈", "보도제목": title, "링크": full_link, "검색어": keyword})
                            found_count += 1
                            
                    except Exception as e_item:
                        print(f"⚠️ [굿타임즈] 기사 처리 중 오류 (내부): {e_item} (키워드: {keyword})")
                        continue
                
                if found_count == 0:
                    results.append({
                        "보도일": f"{days_limit}일 이내 없음",
                        "보도매체": "굿타임즈",
                        "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                        "링크": search_url,
                        "검색어": keyword
                    })
                    
            except TimeoutException:
                print(f"🔍 [굿타임즈] '{keyword}' 검색 결과 없음 (타임아웃, 요소 없음)")
                results.append({
                    "보도일": f"{days_limit}일 이내 없음",
                    "보도매체": "굿타임즈",
                    "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                    "링크": search_url,
                    "검색어": keyword
                })
                
    except Exception as e_page:
        print(f"🔍 [굿타임즈] 테이블 기사 로딩 오류: {e_page}")
    finally:
        if driver:
            try: driver.quit()
            except Exception as e_quit: print(f"⚠️ [굿타임즈] 드라이버 종료 예외: {e_quit}")
    return results

# fetch_hnrsm.py (최적화)
def fetch_hnrsm_multi(keywords, date_limit, days_limit):
    results = []
    base_url = "https://www.hnrsm.com"
    driver = setup_driver(headless=True)
    if not driver: return []
    
    try:
        for keyword in keywords:
            encoded_keyword = quote(keyword)
            search_url = f"{base_url}/news/search.php?q={encoded_keyword}"
            print(f"🌐 [하나로신문] '{keyword}' 검색 중... (URL: {search_url})")

            driver.get(search_url)

            try:
                WebDriverWait(driver, 5).until(
                    EC.presence_of_all_elements_located((By.CSS_SELECTOR, "ul.news_list_skin > li"))
                )
                
                soup = BeautifulSoup(driver.page_source, "html.parser")
                articles = soup.select("ul.news_list_skin > li")
                found_count = 0
                seen_links = set()
                
                for art in articles:
                    try:
                        a_tag = art.select_one("a")
                        title_tag = art.select_one("h3.line_int")
                        date_tag = art.select_one("p.date")
                        
                        if not a_tag or not title_tag or not date_tag: continue
                        
                        title = title_tag.get_text(strip=True)
                        href = a_tag.get("href")
                        link = base_url + href if href and href.startswith("/") else href
                        
                        if not link or link in seen_links: continue
                        seen_links.add(link)
                        
                        date_text = date_tag.get_text(strip=True)
                        pub_date = None
                        
                        for fmt in ("%Y-%m-%d", "%Y.%m.%d", "%Y-%m-%d %H:%M", "%Y.%m.%d %H:%M"):
                            try:
                                pub_date = datetime.strptime(date_text, fmt)
                                break
                            except: continue
                            
                        if pub_date and pub_date >= date_limit:
                            results.append({"보도일": pub_date.strftime("%Y-%m-%d"), "보도매체": "하나로신문", "보도제목": title, "링크": link, "검색어": keyword})
                            found_count += 1
                        
                    except Exception as e_item:
                        print(f"⚠️ [하나로신문] 기사 파싱 오류: {e_item} (키워드: {keyword})")
                        continue

                if found_count == 0:
                    results.append({
                        "보도일": f"{days_limit}일 이내 없음",
                        "보도매체": "하나로신문",
                        "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                        "링크": search_url,
                        "검색어": keyword
                    })
                    
            except TimeoutException:
                print(f"🔍 [하나로신문] '{keyword}' 검색 결과 없음 (타임아웃)")
                results.append({
                    "보도일": f"{days_limit}일 이내 없음",
                    "보도매체": "하나로신문",
                    "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                    "링크": search_url,
                    "검색어": keyword
                })
                
    except Exception as e_page:
        print(f"❌ [하나로신문] 페이지 로드 실패: {e_page}")
    finally:
        if driver:
            try: driver.quit()
            except Exception as e_quit: print(f"⚠️ [하나로신문] 드라이버 종료 예외: {e_quit}")
    return results

# fetch_ehanam.py (최적화)
def fetch_ehanam_multi(keywords, date_limit, days_limit):
    results = []
    base_url = "https://www.ehanam.net"
    driver = setup_driver(headless=True)
    if not driver: return []
    
    try:
        for keyword in keywords:
            search_url = f"{base_url}/news/articleList.html?sc_area=A&view_type=sm&sc_word={quote(keyword)}"
            print(f"🌐 [하남신문] '{keyword}' 검색 중... (URL: {search_url})")

            driver.get(search_url)

            try:
                WebDriverWait(driver, 5).until(
                    EC.presence_of_all_elements_located((By.CSS_SELECTOR, "ul.type2 > li"))
                )
                
                soup = BeautifulSoup(driver.page_source, "html.parser")
                found_count = 0
                
                for li in soup.select("ul.type2 > li"):
                    try:
                        title_tag = li.find("h4", class_="titles")
                        if not title_tag: continue
                        a_tag = title_tag.find("a")
                        if not a_tag: continue
                        
                        title = a_tag.get_text(strip=True)
                        link = a_tag["href"]
                        if not link.startswith("http"):
                            link = base_url + link
                            
                        date_tag_span = li.find("span", class_="byline")
                        if not date_tag_span: continue
                        
                        date_em = date_tag_span.select_one("em:last-child")
                        date_text_full = date_tag_span.get_text()
                        
                        date_match = re.search(r"(\d{4}\.\d{2}\.\d{2})", date_em.get_text(strip=True)) if date_em else re.search(r"(\d{4}\.\d{2}\.\d{2})", date_text_full)
                        
                        if not date_match: continue
                        
                        pub_date = datetime.strptime(date_match.group(1), "%Y.%m.%d")
                        
                        if pub_date >= date_limit:
                            results.append({"보도일": pub_date.strftime("%Y-%m-%d"), "보도매체": "하남신문", "보도제목": title, "링크": link, "검색어": keyword})
                            found_count += 1
                            
                    except Exception as e_item:
                        print(f"⚠️ [하남신문] 개별 파싱 예외: {e_item} (키워드: {keyword})")
                        continue
                
                if found_count == 0:
                    results.append({
                        "보도일": f"{days_limit}일 이내 없음",
                        "보도매체": "하남신문",
                        "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                        "링크": search_url,
                        "검색어": keyword
                    })
                    
            except TimeoutException:
                print(f"🔍 [하남신문] '{keyword}' 검색 결과 없음 (타임아웃)")
                results.append({
                    "보도일": f"{days_limit}일 이내 없음",
                    "보도매체": "하남신문",
                    "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                    "링크": search_url,
                    "검색어": keyword
                })
                
    except Exception as e_page:
           print(f"❌ [하남신문] 페이지 로딩 실패: {e_page}")
    finally:
        if driver:
            try: driver.quit()
            except Exception as e_quit: print(f"⚠️ [하남신문] 드라이버 종료 예외: {e_quit}")
    return results

# fetch_gjilbo.py (최적화)
def fetch_gjilbo_multi(keywords, date_limit, days_limit):
    results = []
    base_url = "http://www.gjilbo.com"
    driver = setup_driver(headless=True)
    if not driver: return []
    
    try:
        for keyword in keywords:
            encoded_keyword = quote(keyword.encode('euc-kr'))
            search_url = f"{base_url}/news/articleList.html?sc_word={encoded_keyword}"
            print(f"🌐 [광주신문] '{keyword}' 검색 중... (URL: {search_url})")

            driver.get(search_url)

            try:
                WebDriverWait(driver, 5).until(
                    EC.presence_of_all_elements_located((By.CSS_SELECTOR, "a[href*='articleView']"))
                )
                
                soup = BeautifulSoup(driver.page_source, "html.parser")
                found_count = 0
                
                for row in soup.select("table tr"):
                    try:
                        a_tag = row.select_one("a[href*='articleView']")
                        if not a_tag: continue
                        
                        title = a_tag.get_text(strip=True)
                        link = base_url + "/news/" + a_tag["href"]
                        
                        tds = row.select("td")
                        if not tds: continue
                        
                        date_tag = row.select_one("td.View_Sm_Date")
                        if not date_tag:
                            date_tag = tds[-1]
                        
                        date_text = date_tag.get_text(strip=True).split(" ")[0]
                        
                        if not re.match(r"\d{4}-\d{2}-\d{2}", date_text): continue
                        
                        pub_date = datetime.strptime(date_text, "%Y-%m-%d")
                        
                        if pub_date >= date_limit and keyword in title:
                            results.append({
                                "보도일": pub_date.strftime("%Y-%m-%d"),
                                "보도매체": "광주신문",
                                "보도제목": title,
                                "링크": link,
                                "검색어": keyword
                            })
                            found_count += 1
                            
                    except Exception as e_item:
                        print(f"⚠️ [광주신문] 예외: {e_item} (키워드: {keyword})")
                        continue

                if found_count == 0:
                    results.append({
                        "보도일": f"{days_limit}일 이내 없음",
                        "보도매체": "광주신문",
                        "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                        "링크": search_url,
                        "검색어": keyword
                    })
                    
            except TimeoutException:
                print(f"🔍 [광주신문] '{keyword}' 검색 결과 없음 (타임아웃)")
                results.append({
                    "보도일": f"{days_limit}일 이내 없음",
                    "보도매체": "광주신문",
                    "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                    "링크": search_url,
                    "검색어": keyword
                })
                
    except Exception as e_page:
        print(f"🔍 [광주신문] 검색 중 오류: {e_page}")
    finally:
        if driver:
            try: driver.quit()
            except Exception as e_quit: print(f"⚠️ [광주신문] 드라이버 종료 예외: {e_quit}")
    return results

# fetch_jungbusisa.py (최적화)
def fetch_jungbusisa_multi(keywords, date_limit, days_limit):
    results = []
    base_url = "https://www.gninews.co.kr"
    driver = setup_driver(headless=True)
    if not driver: return []
    
    try:
        for keyword in keywords:
            encoded_keyword = quote(keyword)
            search_url = f"{base_url}/news/search_result.html?search={encoded_keyword}"
            print(f"🌐 [중부시사] '{keyword}' 검색 중... (URL: {search_url})")

            driver.get(search_url)

            try:
                WebDriverWait(driver, 5).until(
                    EC.presence_of_all_elements_located((By.CSS_SELECTOR, "ul.art_list_all > li"))
                )
                
                soup = BeautifulSoup(driver.page_source, "html.parser")
                articles = soup.select("ul.art_list_all > li")
                found_count = 0
                
                for article in articles:
                    try:
                        title_tag = article.select_one("h2.cmp.c2")
                        link_tag = article.select_one("a[href*='article.html']")
                        date_tag = article.select_one("li.date")
                        
                        if not (title_tag and link_tag and date_tag): continue
                        
                        title = title_tag.get_text(strip=True)
                        href = link_tag["href"]
                        full_link = urljoin(base_url, '/news/' + href.lstrip('./'))
                        date_text = date_tag.get_text(strip=True)
                        
                        pub_date = datetime.strptime(date_text.split()[0], "%Y.%m.%d")
                        
                        if pub_date >= date_limit:
                            results.append({"보도일": pub_date.strftime("%Y-%m-%d"), "보도매체": "중부시사신문", "보도제목": title, "링크": full_link, "검색어": keyword})
                            found_count += 1
                            
                    except Exception as e_item:
                        print(f"⚠️ [중부시사] 파싱 예외: {e_item} (키워드: {keyword})")
                        continue
                
                if found_count == 0:
                    results.append({
                        "보도일": f"{days_limit}일 이내 없음",
                        "보도매체": "중부시사신문",
                        "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                        "링크": search_url,
                        "검색어": keyword
                    })
                    
            except TimeoutException:
                print(f"🔍 [중부시사] '{keyword}' 검색 결과 없음 (타임아웃)")
                results.append({
                    "보도일": f"{days_limit}일 이내 없음",
                    "보도매체": "중부시사신문",
                    "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                    "링크": search_url,
                    "검색어": keyword
                })
                
    except Exception as e_page:
        print(f"❌ [중부시사] 기사 로딩 실패: {e_page}")
    finally:
        if driver:
            try: driver.quit()
            except Exception as e_quit: print(f"⚠️ [중부시사] 드라이버 종료 예외: {e_quit}")
    return results

# fetch_portalnews.py (최적화)
def fetch_portalnews_multi(keywords, date_limit, days_limit):
    results = []
    base_url = "https://www.portalnews.co.kr"
    driver = setup_driver(headless=True)
    if not driver: return []
    
    try:
        for keyword in keywords:
            search_url = f"{base_url}/news/search_result.html?search_mode=multi&s_title=1&s_body=1&search={quote(keyword)}&s_sdate=&s_edate=&catno=106"
            print(f"🌐 [포탈뉴스] '{keyword}' 검색 중... (URL: {search_url})")

            driver.get(search_url)

            try:
                WebDriverWait(driver, 5).until(
                    EC.presence_of_all_elements_located((By.CSS_SELECTOR, "ul.art_list_all li"))
                )
                
                soup = BeautifulSoup(driver.page_source, "html.parser")
                items = soup.select("ul.art_list_all li")
                found_count = 0
                no_date_count = 0 
                
                for item in items:
                    try:
                        title_tag = item.select_one("h2")
                        link_tag = item.select_one("a")
                        date_tag = item.select_one("li.date")
                        
                        if not title_tag or not link_tag: continue
                        
                        title = title_tag.text.strip()
                        link = link_tag.get("href", "")
                        
                        if not link.startswith("http"):
                            link = base_url + link if link.startswith("/") else base_url + "/news/" + link
                            
                        date_str = "기사에서 직접 확인 바람"
                        pub_date = None
                        
                        if date_tag:
                            date_text = date_tag.text.strip()
                            try:
                                pub_date = datetime.strptime(date_text.split()[0], "%Y.%m.%d")
                                date_str = pub_date.strftime("%Y-%m-%d")
                            except:
                                pass 
                        
                        if pub_date and pub_date < date_limit: continue
                        
                        if not pub_date:
                            no_date_count += 1
                            if no_date_count > 5 and len(keywords) > 1: continue 
                        
                        # [v5.0] '기타지방지'는 날짜 불명확 시 '기사에서 직접 확인 바람'으로 표기
                        results.append({"보도일": date_str, "보도매체": "포탈뉴스통신", "보도제목": title, "링크": link, "검색어": keyword})
                        found_count += 1
                        
                    except Exception as e_item:
                        print(f"⚠️ [포탈뉴스] 기사 처리 오류: {e_item} (키워드: {keyword})")
                        continue

                if found_count == 0:
                    results.append({
                        "보도일": f"{days_limit}일 이내 없음",
                        "보도매체": "포탈뉴스통신",
                        "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                        "링크": search_url,
                        "검색어": keyword
                    })
                    
            except TimeoutException:
                print(f"🔍 [포탈뉴스] '{keyword}' 검색 결과 없음 (타임아웃)")
                results.append({
                    "보도일": f"{days_limit}일 이내 없음",
                    "보도매체": "포탈뉴스통신",
                    "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                    "링크": search_url,
                    "검색어": keyword
                })
                
    except Exception as e_page:
        print(f"❌ [포탈뉴스] 페이지 로딩 오류: {e_page}")
    finally:
        if driver:
            try: driver.quit()
            except Exception as e_quit: print(f"⚠️ [포탈뉴스] 드라이버 종료 예외: {e_quit}")
    return results

# fetch_kgyonhap.py (최적화)
def fetch_kgyonhap_multi(keywords, date_limit, days_limit):
    results = []
    base_url = "https://kgyonhapnews.net"
    driver = setup_driver(headless=True)
    if not driver: return []
    
    try:
        for keyword in keywords:
            encoded_keyword = quote(keyword)
            search_url = f"{base_url}/search.html?submit=submit&search_and=1&search_exec=all&search_section=all&news_order=1&search={encoded_keyword}&imageField.x=0&imageField.y=0"
            print(f"🌐 [경기연합] '{keyword}' 검색 중... (URL: {search_url})")

            driver.get(search_url)

            try:
                WebDriverWait(driver, 5).until(
                    EC.presence_of_all_elements_located((By.CSS_SELECTOR, "div.search_result_list_box dl"))
                )
                
                soup = BeautifulSoup(driver.page_source, "html.parser")
                found_count = 0
                
                for dl in soup.select("div.search_result_list_box dl"):
                    try:
                        a_tag = dl.select_one("dt > a")
                        if not a_tag: continue
                        
                        title = a_tag.get_text(strip=True)
                        link = a_tag["href"]
                        if not link.startswith("http"):
                            link = base_url + link
                            
                        date_tag = dl.select_one("dd.etc")
                        if not date_tag: continue
                        
                        date_match = re.search(r"\d{4}\.\d{2}\.\d{2}", date_tag.get_text())
                        if not date_match: continue
                        
                        pub_date = datetime.strptime(date_match.group(), "%Y.%m.%d")
                        
                        if pub_date >= date_limit:
                            results.append({"보도일": pub_date.strftime("%Y-%m-%d"), "보도매체": "경기연합뉴스", "보도제목": title, "링크": link, "검색어": keyword})
                            found_count += 1
                            
                    except Exception as e_item:
                        print(f"⚠️ [경기연합] 예외: {e_item} (키워드: {keyword})")
                        continue

                if found_count == 0:
                    results.append({
                        "보도일": f"{days_limit}일 이내 없음",
                        "보도매체": "경기연합뉴스",
                        "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                        "링크": search_url,
                        "검색어": keyword
                    })
                    
            except TimeoutException:
                print(f"🔍 [경기연합] '{keyword}' 검색 결과 없음 (타임아웃)")
                results.append({
                    "보도일": f"{days_limit}일 이내 없음",
                    "보도매체": "경기연합뉴스",
                    "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                    "링크": search_url,
                    "검색어": keyword
                })
                
    except Exception as e_page:
        print(f"❌ [경기연합] 페이지 로딩 오류: {e_page}")
    finally:
        if driver:
            try: driver.quit()
            except Exception as e_quit: print(f"⚠️ [경기연합] 드라이버 종료 예외: {e_quit}")
    return results

# fetch_kocus.py (최적화)
def fetch_kocus_multi(keywords, date_limit, days_limit):
    results = []
    base_url = "http://www.kocus.com"
    search_path = "/news/articleList.html"
    driver = setup_driver(headless=True)
    if not driver: return []
    
    try:
        for keyword in keywords:
            params = {
                "sc_area": "A", "sc_word": keyword, "view_type": ""
            }
            # EUC-KR 인코딩이 필요함
            full_url = f"{base_url}{search_path}?{urlencode(params, encoding='euc-kr')}"
            print(f"🌐 [교차로저널] '{keyword}' 검색 중... (URL: {full_url})")

            driver.get(full_url)

            try:
                WebDriverWait(driver, 5).until(
                    EC.presence_of_all_elements_located((By.CSS_SELECTOR, "td.list-titles"))
                )
                
                soup = BeautifulSoup(driver.page_source, "html.parser")
                title_links = soup.select("td.list-titles a")
                found_count = 0
                no_date_count = 0
                
                for link_tag in title_links:
                    try:
                        title = link_tag.text.strip()
                        article_url = urljoin(full_url, link_tag.get("href", ""))
                        row = link_tag.find_parent("tr")
                        date_cell = row.find("td", class_="list-times") if row else None
                        date_text = date_cell.text.strip() if date_cell else ""
                        
                        date_str = "기사에서 직접 확인 바람"
                        pub_date = None
                        
                        try:
                            pub_date = datetime.strptime(date_text, "%Y-%m-%d %H:%M")
                        except Exception:
                            try:
                                pub_date = datetime.strptime(date_text, "%Y.%m.%d")
                            except Exception:
                                pass
                                
                        if pub_date:
                            if pub_date < date_limit: continue
                            date_str = pub_date.strftime("%Y-%m-%d")
                        else:
                            no_date_count += 1
                            if no_date_count > 5 and len(keywords) > 1: continue 

                        results.append({"보도일": date_str, "보도매체": "교차로저널", "보도제목": title, "링크": article_url, "검색어": keyword})
                        found_count += 1
                        
                    except Exception as e_item:
                        print(f"⚠️ [교차로저널] 기사 처리 오류: {e_item} (키워드: {keyword})")
                        continue

                if found_count == 0:
                    results.append({
                        "보도일": f"{days_limit}일 이내 없음",
                        "보도매체": "교차로저널",
                        "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                        "링크": full_url,
                        "검색어": keyword
                    })
                    
            except TimeoutException:
                print(f"🔍 [교차로저널] '{keyword}' 검색 결과 없음 (타임아웃)")
                results.append({
                    "보도일": f"{days_limit}일 이내 없음",
                    "보도매체": "교차로저널",
                    "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                    "링크": full_url,
                    "검색어": keyword
                })
                
    except Exception as e_page:
        print(f"⚠️ [교차로저널] 페이지 로드 실패: {e_page}")
    finally:
        if driver:
            try: driver.quit()
            except Exception as e_quit: print(f"⚠️ [교차로저널] 드라이버 종료 예외: {e_quit}")
    return results

# fetch_vision21.py (최적화)
def fetch_vision21_multi(keywords, date_limit, days_limit):
    results = []
    base_url = "https.www.vision21.kr"
    driver = setup_driver(headless=True)
    if not driver: return []
    
    try:
        for keyword in keywords:
            params = {"search": keyword}
            search_url = f"{base_url}/news/search_result.html?{urlencode(params, encoding='utf-8')}"
            print(f"🌐 [비전21뉴스] '{keyword}' 검색 중... (URL: {search_url})")

            driver.get(search_url)
            
            try:
                try:
                    WebDriverWait(driver, 1).until(EC.alert_is_present())
                    alert = driver.switch_to.alert
                    alert_text = alert.text
                    alert.accept()
                    print(f"❌ [비전21뉴스] '{keyword}' 검색 얼럿 발생: {alert_text}")
                    if "검색결과가 없습니다" in alert_text:
                        raise TimeoutException("검색 결과 없음 (Alert)")
                except TimeoutException as e:
                    if "Alert" in str(e): raise 
                    pass 

                WebDriverWait(driver, 5).until(
                    EC.presence_of_all_elements_located((By.CSS_SELECTOR, "ul.art_list_all li"))
                )
                
                soup = BeautifulSoup(driver.page_source, "html.parser")
                articles = soup.select("ul.art_list_all > li")
                found_count = 0
                no_date_count = 0
                
                for article in articles:
                    try:
                        title_tag = article.select_one("h2.cmp.c2") or article.select_one("h4.cmp.c2")
                        link_tag = article.select_one("a[href*='article.html']")
                        date_tag = article.select_one("li.date")
                        
                        if not title_tag or not link_tag: continue
                        
                        title = title_tag.get_text(strip=True)
                        href = link_tag.get("href", "")
                        article_url = urljoin(base_url + "/news/", href)
                        date_text = date_tag.get_text(strip=True) if date_tag else ""
                        
                        date_str = "기사에서 직접 확인 바람"
                        pub_date = None
                        
                        if date_text:
                            date_str_clean = date_text.split()[0].replace('-', '.')
                            try:
                                pub_date = datetime.strptime(date_str_clean, "%Y.%m.%d")
                                date_str = pub_date.strftime("%Y-%m-%d")
                            except Exception as pe:
                                pass

                        if pub_date and pub_date < date_limit: continue
                        
                        if not pub_date:
                            no_date_count += 1
                            if no_date_count > 5 and len(keywords) > 1: continue 

                        results.append({"보도일": date_str, "보도매체": "비전21뉴스", "보도제목": title, "링크": article_url, "검색어": keyword})
                        found_count += 1
                        
                    except Exception as e_item:
                        print(f"⚠️ [비전21뉴스] 기사 처리 오류: {e_item} (키워드: {keyword})")
                        continue

                if found_count == 0:
                    results.append({
                        "보도일": f"{days_limit}일 이내 없음",
                        "보도매체": "비전21뉴스",
                        "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                        "링크": search_url,
                        "검색어": keyword
                    })
                    
            except TimeoutException:
                print(f"🔍 [비전21뉴스] '{keyword}' 검색 결과 없음 (타임아웃/Alert)")
                results.append({
                    "보도일": f"{days_limit}일 이내 없음",
                    "보도매체": "비전21뉴스",
                    "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                    "링크": search_url,
                    "검색어": keyword
                })
                
    except Exception as e_page:
        print(f"❌ [비전21뉴스] 전체 페이지 처리 오류: {e_page}")
    finally:
        if driver:
            try: driver.quit()
            except Exception as e_quit: print(f"⚠️ [비전21뉴스] 드라이버 종료 예외: {e_quit}")
    return results

# fetch_newstoday24.py (최적화)
def fetch_newstoday24_multi(keywords, date_limit, days_limit):
    results = []
    base_url = "https://www.newstoday.or.kr"
    driver = setup_driver(headless=True)
    if not driver: return []
    
    try:
        for keyword in keywords:
            search_url = f"{base_url}/news/articleList.html?sc_area=A&view_type=sm&sc_word={quote(keyword)}"
            print(f"🌐 [뉴스투데이24] '{keyword}' 검색 중... (URL: {search_url})")

            driver.get(search_url)

            try:
                WebDriverWait(driver, 5).until(
                    EC.presence_of_all_elements_located((By.CSS_SELECTOR, "div.view-cont"))
                )
                
                soup = BeautifulSoup(driver.page_source, "html.parser")
                article_blocks = soup.select("div.view-cont")
                found_count = 0
                
                for block in article_blocks:
                    try:
                        title_tag = block.select_one("h4.titles a")
                        date_tag = block.select_one("em.replace-date")
                        
                        if not title_tag or not date_tag: continue
                        
                        title = title_tag.get_text(strip=True)
                        href = title_tag.get("href", "")
                        link = f"{base_url}{href}" if href.startswith("/") else href
                        date_text_raw = date_tag.get_text(strip=True)
                        
                        if not date_text_raw: continue
                        date_text = date_text_raw.split(" ")[0]
                        
                        pub_date = datetime.strptime(date_text, "%Y.%m.%d")
                        
                        if pub_date >= date_limit:
                            results.append({"보도일": pub_date.strftime("%Y-%m-%d"), "보도매체": "뉴스투데이24", "보도제목": title, "링크": link, "검색어": keyword})
                            found_count += 1
                            
                    except Exception as e_item:
                        print(f"⚠️ [뉴스투데이24] 파싱 예외: {e_item} (키워드: {keyword})")
                        continue

                if found_count == 0:
                    results.append({
                        "보도일": f"{days_limit}일 이내 없음",
                        "보도매체": "뉴스투데이24",
                        "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                        "링크": search_url,
                        "검색어": keyword
                    })
                    
            except TimeoutException:
                print(f"🔍 [뉴스투데이24] '{keyword}' 검색 결과 없음 (타임아웃)")
                results.append({
                    "보도일": f"{days_limit}일 이내 없음",
                    "보도매체": "뉴스투데이24",
                    "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                    "링크": search_url,
                    "검색어": keyword
                })
                
    except Exception as e_page:
        print(f"🔍 [뉴스투데이24] 기사 로딩 오류: {e_page}")
    finally:
        if driver:
            try: driver.quit()
            except Exception as e_quit: print(f"⚠️ [뉴스투데이24] 드라이버 종료 예외: {e_quit}")
    return results

# fetch_tgh.py (최적화)
def fetch_tgh_multi(keywords, date_limit, days_limit):
    results = []
    base_url = "http://www.tgh.kr"
    search_path = "/news/articleList.html"
    driver = setup_driver_compatible(headless=True)
    if not driver: return []
    
    try:
        for keyword in keywords:
            params = {
                "sc_area": "A", "sc_word": keyword, "sc_order_by": "E", "view_type": ""
            }
            full_url = f"{base_url}{search_path}?{urlencode(params, encoding='euc-kr')}"
            print(f"🌐 [투데이광주하남] '{keyword}' 검색 중... (URL: {full_url})")

            driver.get(full_url)

            try:
                WebDriverWait(driver, 5).until(
                    EC.presence_of_all_elements_located((By.CSS_SELECTOR, "td.list-titles a"))
                )
                
                soup = BeautifulSoup(driver.page_source, "html.parser")
                title_links = soup.select("td.list-titles a")
                found_count = 0
                no_date_count = 0
                
                for link_tag in title_links:
                    try:
                        title = link_tag.get_text(strip=True)
                        article_url = urljoin(full_url, link_tag.get("href", ""))
                        row = link_tag.find_parent("tr")
                        date_cell = row.find("td", class_="list-times") if row else None
                        
                        if not date_cell: continue
                        
                        date_text = date_cell.get_text(strip=True)
                        date_str = "기사에서 직접 확인 바람"
                        pub_date = None
                        
                        if date_text.strip():
                            try:
                                pub_date = datetime.strptime(date_text, "%Y-%m-%d %H:%M")
                                date_str = pub_date.strftime("%Y-%m-%d")
                            except Exception:
                                pass
                                
                        if pub_date and pub_date < date_limit: continue
                        
                        if not pub_date:
                            no_date_count += 1
                            if no_date_count > 5 and len(keywords) > 1: continue 

                        results.append({"보도일": date_str, "보도매체": "투데이광주하남", "보도제목": title, "링크": article_url, "검색어": keyword})
                        found_count += 1
                        
                    except Exception as e_item:
                        print(f"⚠️ [투데이광주하남] 기사 처리 오류: {e_item} (키워드: {keyword})")
                        continue

                if found_count == 0:
                    results.append({
                        "보도일": f"{days_limit}일 이내 없음",
                        "보도매체": "투데이광주하남",
                        "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다. (검색어 링크 확인)",
                        "링크": full_url,
                        "검색어": keyword
                    })
                    
            except TimeoutException:
                print(f"🔍 [투데이광주하남] '{keyword}' 검색 결과 없음 (타임아웃)")
                results.append({
                    "보도일": f"{days_limit}일 이내 없음",
                    "보도매체": "투데이광주하남",
                    "보도제목": f"웹사이트 응답 지연으로 검색 실패 (키워드: '{keyword}')",
                    "링크": full_url,
                    "검색어": keyword
                })
                
    except Exception as e_page:
        print(f"⚠️ [투데이광주하남] 페이지 로드 실패: {e_page}")
    finally:
        if driver:
            try: driver.quit()
            except Exception as e_quit: print(f"⚠️ [투데이광주하남] 드라이버 종료 예외: {e_quit}")
    return results

# fetch_mediareport.py (최적화)
def fetch_mediareport_multi(keywords, date_limit, days_limit):
    results = []
    base_url = "https://mediareport.co.kr"
    driver = setup_driver(headless=True)
    if not driver: return []
    
    try:
        for keyword in keywords:
            encoded_keyword = quote(keyword)
            search_url = f"{base_url}/search.html?submit=submit&search={encoded_keyword}&search_and=2&search_exec=all&search_section=all&news_order=1"
            print(f"🌐 [미디어리포트] '{keyword}' 검색 중... (URL: {search_url})")

            driver.get(search_url)

            try:
                WebDriverWait(driver, 5).until(
                    EC.presence_of_all_elements_located((By.CSS_SELECTOR, "div.search_result_list_box dl"))
                )
                
                soup = BeautifulSoup(driver.page_source, "html.parser")
                found_count = 0
                
                for dl_tag in soup.select("div.search_result_list_box dl"):
                    try:
                        dt_tag = dl_tag.find("dt")
                        if not dt_tag: continue
                        a_tag = dt_tag.find("a")
                        if not a_tag: continue
                        
                        title = a_tag.get_text(strip=True)
                        link = a_tag["href"]
                        if not link.startswith("http"):
                            link = base_url + link
                            
                        date_tag = dl_tag.find("dd", class_="etc")
                        if not date_tag: continue
                        
                        date_match = re.search(r"\d{4}\.\d{2}\.\d{2}", date_tag.get_text())
                        if not date_match: continue
                        
                        pub_date = datetime.strptime(date_match.group(), "%Y.%m.%d")
                        
                        if pub_date >= date_limit:
                            results.append({"보도일": pub_date.strftime("%Y-%m-%d"), "보도매체": "미디어리포트", "보도제목": title, "링크": link, "검색어": keyword})
                            found_count += 1
                            
                    except Exception as e_item:
                        print(f"⚠️ [미디어리포트] 예외: {e_item} (키워드: {keyword})")
                        continue

                if found_count == 0:
                    results.append({
                        "보도일": f"{days_limit}일 이내 없음",
                        "보도매체": "미디어리포트",
                        "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                        "링크": search_url,
                        "검색어": keyword
                    })
                    
            except TimeoutException:
                print(f"🔍 [미디어리포트] '{keyword}' 검색 결과 없음 (타임아웃)")
                results.append({
                    "보도일": f"{days_limit}일 이내 없음",
                    "보도매체": "미디어리포트",
                    "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                    "링크": search_url,
                    "검색어": keyword
                })
                
    except Exception as e_page:
        print(f"❌ [미디어리포트] 페이지 로딩 오류: {e_page}")
    finally:
        if driver:
            try: driver.quit()
            except Exception as e_quit: print(f"⚠️ [미디어리포트] 드라이버 종료 예외: {e_quit}")
    return results

# fetch_sudokwon.py (최적화)
def fetch_sudokwon_multi(keywords, date_limit, days_limit):
    results = []
    base_url = "http://www.sudokwon.com"
    driver = setup_driver_compatible(headless=True)
    if not driver: return []
    
    try:
        for keyword in keywords:
            encoded_keyword = quote(keyword.encode('euc-kr'))
            search_url = f"{base_url}/searchs.php?searchword={encoded_keyword}&x=15&y=20"
            print(f"🌐 [수도권일보] '{keyword}' 검색 중... (URL: {search_url})")

            driver.get(search_url)

            try:
                # ✅ [수정] 대기 시간을 5초에서 10초로 늘렸습니다.
                WebDriverWait(driver, 10).until(
                    EC.presence_of_all_elements_located((By.CSS_SELECTOR, "a.sublist"))
                )
                
                soup = BeautifulSoup(driver.page_source, "html.parser")
                articles = soup.select("td[style*='padding:10px']")
                found_count = 0
                
                for article_td in articles:
                    try:
                        title_tag = article_td.select_one("a.sublist")
                        date_tag = article_td.select_one("span.date")
                        
                        if not (title_tag and date_tag): continue
                        
                        title = title_tag.get_text(strip=True)
                        href = title_tag.get("href", "")
                        full_link = f"{base_url}/{href.lstrip('/')}"
                        date_text = date_tag.get_text(strip=True)
                        
                        if re.match(r"\d{4}\.\s?\d{2}\.\s?\d{2}", date_text):
                            pub_date = datetime.strptime(date_text.replace(" ", ""), "%Y.%m.%d")
                            if pub_date >= date_limit:
                                results.append({"보도일": pub_date.strftime("%Y-%m-%d"), "보도매체": "수도권일보", "보도제목": title, "링크": full_link, "검색어": keyword})
                                found_count += 1
                        
                    except Exception as e_item:
                        print(f"⚠️ [수도권일보] 파싱 예외: {e_item} (키워드: {keyword})")
                        continue

                if found_count == 0:
                    results.append({
                        "보도일": f"{days_limit}일 이내 없음",
                        "보도매체": "수도권일보",
                        "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다. (검색어 링크 확인)",
                        "링크": search_url,
                        "검색어": keyword
                    })
                    
            except TimeoutException:
                print(f"🔍 [수도권일보] '{keyword}' 검색 결과 없음 (타임아웃 - 10초 대기)")
                results.append({
                    "보도일": f"{days_limit}일 이내 없음",
                    "보도매체": "수도권일보",
                    "보도제목": f"웹사이트 응답 지연으로 검색 실패 (키워드: '{keyword}')",
                    "링크": search_url,
                    "검색어": keyword
                })
                
    except Exception as e_page:
        print(f"❌ [수도권일보] 기사 로딩 실패: {e_page}")
    finally:
        if driver:
            try: driver.quit()
            except Exception as e_quit: print(f"⚠️ [수도권일보] 드라이버 종료 예외: {e_quit}")
    return results

# =============================================================================
# 6. [v5.0] Fetcher 함수 그룹 3 (Google 뉴스)
# (이하 v5.0 원본과 동일 ... fetch_google_news)
# =============================================================================

def shorten_google_url(url):
    """구글 리디렉션 URL에서 실제 URL을 추출합니다."""
    if "google.com/url?" in url:
        try:
            # URL의 쿼리 문자열을 파싱하여 'url' 파라미터 값을 추출
            parsed_url = urlparse(url)
            query_params = parse_qs(parsed_url.query)
            if 'url' in query_params:
                return query_params['url'][0]
        except Exception as e:
            print(f"URL 단축 중 오류 발생: {e} (URL: {url})")
            return url # 오류 발생 시 원본 URL 반환
    # 구글 리디렉션 URL이 아니면 원본 URL 반환
    return url

def fetch_google_news_feed(keywords, date_limit, days_limit):
    """
    Google News RSS 피드를 사용하여 여러 키워드에 대한 뉴스를 검색합니다.
    ✅ [v5.0.3 수정] Selenium을 사용하여 최종 URL을 확인합니다. (느릴 수 있음)
    """
    source_name = "Google뉴스"
    results = []
    
    found_for_keyword = {keyword: False for keyword in keywords}

    driver = None
    try:
        # 호환성 드라이버(리소스 차단 없음)를 사용해야 리디렉션이 잘 됩니다.
        driver = setup_driver_compatible(headless=True)
        if not driver:
            print("‼️ [Google뉴스] Selenium 드라이버를 시작할 수 없어 URL 단축을 건너뜁니다.")
    except Exception as e:
        print(f"‼️ [Google뉴스] Selenium 드라이버 시작 중 오류: {e}")

    for keyword in keywords:
        # Google News RSS 피드 URL 생성
        # ✅ [v5.0.1] '광주하남교육지원청' OR '광주하남' -> "광주하남교육지원청" 통일
        search_keyword = keyword
        if keyword == "광주하남":
            search_keyword = "광주하남교육지원청"
        
        # ✅ [v5.0.1] 정확한 검색을 위해 키워드에 큰따옴표 추가
        query = f'"{search_keyword}"'
        
        # ✅ [v5.0.1] daterange 대신 after/before 사용 (더 안정적)
        # ex: "광주하남교육지원청" after:2024-05-20 before:2024-05-27
        end_date = datetime.now()
        start_date = end_date - timedelta(days=days_limit)
        query += f" after:{start_date.strftime('%Y-%m-%d')} before:{end_date.strftime('%Y-%m-%d')}"

        # RSS 피드 URL
        url = f"https://news.google.com/rss/search?q={quote(query)}&hl=ko&gl=KR&ceid=KR:ko"
        print(f"🌐 [{source_name}] '{keyword}' 검색 중... (URL: {url})")

        try:
            feed = feedparser.parse(url)
            
            if not feed.entries:
                # ✅ [v5.0.1] '없음' 메시지는 마지막에 한번만 추가
                continue

            for entry in feed.entries:
                try:
                    # --- 날짜 파싱 ---
                    pub_date_parsed = entry.get("published_parsed")
                    if not pub_date_parsed:
                        continue
                    
                    pub_date = datetime.fromtimestamp(time.mktime(pub_date_parsed))
                    
                    # --- 기간 필터링 ---
                    if pub_date < date_limit:
                        continue

                    # --- 제목 및 링크 파싱 ---
                    title = entry.title
                    original_link = entry.link
                    final_link = original_link

                    # ✅ [v5.0.3 수정] Selenium으로 링크 단축
                    if driver and "news.google.com/rss/articles/" in original_link:
                        try:
                            driver.set_page_load_timeout(15) # 타임아웃 설정
                            driver.get(original_link)
                            # URL이 바뀌고, 더 이상 구글이 아닐 때까지 최대 10초 대기
                            WebDriverWait(driver, 10).until(
                                lambda d: d.current_url != original_link and "google.com" not in d.current_url
                            )
                            final_link = driver.current_url
                            print(f"  [링크 단축] 성공: {final_link[:70]}...")
                        except TimeoutException:
                            final_link = driver.current_url 
                            print(f"  [링크 단축] 타임아웃. 현재 URL 사용: {final_link[:70]}...")
                        except Exception as e:
                            print(f"  [링크 단축] 오류: {e}")
                            final_link = original_link # 오류 시 원본 링크 사용
                    else:
                        # Selenium 드라이버가 없거나 일반 링크일 경우 기존 로직 사용
                        final_link = shorten_google_url(original_link)


                    # --- 결과 추가 ---
                    results.append({
                        "보도일": pub_date.strftime("%Y-%m-%d"),
                        "보도매체": source_name,
                        "보도제목": title,
                        "링크": final_link,
                        "검색어": keyword
                    })
                    found_for_keyword[keyword] = True

                except Exception as e:
                    print(f"[{source_name}] 개별 기사 처리 중 오류: {e} (제목: {entry.get('title', 'N/A')})")
                    continue
        
        except Exception as e:
            print(f"[{source_name}] '{keyword}' 피드 파싱 중 오류: {e}")
            # ✅ [v5.0.1] '없음' 메시지는 마지막에 한번만 추가
            continue

    # ✅ [v5.0.3 수정] 드라이버 종료
    if driver:
        driver.quit()

    # ✅ [v5.0.1] 검색 결과가 없는 키워드에 대해 '없음' 메시지 추가
    for keyword, found in found_for_keyword.items():
        if not found:
            # '없음' 메시지에 대한 URL은 실제 검색했던 RSS URL을 제공
            search_keyword = keyword
            if keyword == "광주하남":
                search_keyword = "광주하남교육지원청"
            query = f'"{search_keyword}"'
            end_date = datetime.now()
            start_date = end_date - timedelta(days=days_limit)
            query += f" after:{start_date.strftime('%Y-%m-%d')} before:{end_date.strftime('%Y-%m-%d')}"
            url = f"https://news.google.com/rss/search?q={quote(query)}&hl=ko&gl=KR&ceid=KR:ko"
            
            results.append({
                "보도일": "없음",
                "보도매체": source_name,
                "보도제목": f"최근 {days_limit}일 이내 '{keyword}' 관련 기사가 없습니다.",
                "링크": url,
                "검색어": keyword
            })
            
    return results

# =============================================================================
# 7. [v5.0.1 수정] Tkinter GUI 애플리케이션 (통합 및 수정)
# =============================================================================

class NewsScraperApp:
    def __init__(self, root):
        self.root = root
        # ✅ [v5.0.1] GUI 타이틀 변경 (보도자료 뷰어 기능 통합)
        self.root.title("통합 뉴스 스크레이퍼 v5.0.1 (보도자료 엑셀 저장 기능 포함)") 
        
        # [v5.0] GUI 크기 (16대 지방지 기준 1125x750) -> 1125x800 (체크박스 공간)
        self.root.geometry("1125x800") 

        self.goegh_base_url = "https://www.goegh.kr/goegh/na/ntt/selectNttList.do?mi=8686&bbsId=5041"
        self.goegh_page1_url = self.goegh_base_url + "&pageIndex=1"

        self.earliest_date_clicked = None

        # ✅ [v5.0.1 신규] 엑셀 저장을 위한 원본 데이터 저장소 (v1.4 기준)
        self.tree_data = {}
        # ✅ [Word 저장] 최신 통합 검색 결과 캐시
        self.latest_articles_df = None

        self.style = ttk.Style()
        self.style.configure('.', font=('Malgun Gothic', 12))
        self.style.configure('TLabelframe.Label', font=('Malgun Gothic', 12, 'bold'))
        self.style.configure('Status.TLabel', font=('Malgun Gothic', 11))
        self.style.configure('Treeview.Heading', font=('Malgun Gothic', 11, 'bold'))
        # [v5.0] '6대 지방지'의 Treeview 폰트(12) 및 rowheight 적용
        self.style.configure('Treeview', font=('Malgun Gothic', 12), rowheight=int(12 * 2.2))

        # ✅ (v1.4) [요청 3] 강조 버튼 스타일 추가
        self.style.configure('Emph.TButton', font=('Malgun Gothic', 12, 'bold'), foreground='#0000AA')

        main_frame = ttk.Frame(root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # --- 1. [참고] 프레임 ---
        goegh_frame = ttk.LabelFrame(main_frame, text=" [참고] 광주하남교육지원청 보도자료 ", padding="10")
        goegh_frame.pack(fill=tk.X, expand=False, pady=(0, 5))
        link_frame = ttk.Frame(goegh_frame)
        link_frame.pack(fill=tk.X)
        link_label = ttk.Label(link_frame, text="▶ 보도배포(성과팀->언론) 링크 (클릭)",
                               foreground="blue", cursor="hand2", font=('Malgun Gothic', 12, 'underline'))
        link_label.pack(side=tk.LEFT, padx=5)
        link_label.bind("<Button-1>", self.open_link)

        # --- 2. [Step 1] 보도자료 목록 프레임 (v5.0.1 수정) ---
        press_release_frame = ttk.LabelFrame(main_frame, text=" [Step 1] 보도자료 목록 확인 (더블클릭 시 검색어 누적) ", padding="10")
        press_release_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        # ✅ [v5.0.1 수정] v1.4의 버튼 프레임 레이아웃 적용
        button_frame = ttk.Frame(press_release_frame)
        button_frame.pack(fill=tk.X, pady=(0, 10))

        # (v1.4) [수정] (1) 전체 목록 가져오기 - command 변경
        self.fetch_press_button = ttk.Button(button_frame, text="보도배포 목록 가져오기(1페이지)", command=self.start_fetch_press_releases_all)
        self.fetch_press_button.pack(side=tk.LEFT, ipady=3, padx=(0, 10))
        
        # ✅ (v1.4) [신규] (2) 오늘 날짜 버튼 추가
        today_str = datetime.now().strftime('%Y-%m-%d')
        self.fetch_today_button = ttk.Button(
            button_frame, 
            text=f"'{today_str}' 일자 목록 가져오기", 
            command=self.start_fetch_press_releases_today,
            style='Emph.TButton' # 강조 스타일 적용
        )
        self.fetch_today_button.pack(side=tk.LEFT, ipady=3, padx=(5, 10))

        # ✅ (v1.4) [신규] (3) 엑셀 다운로드 버튼 추가
        self.export_button = ttk.Button(button_frame, text="엑셀파일로 다운", command=self.export_to_excel)
        self.export_button.pack(side=tk.LEFT, ipady=3, padx=(10, 0))
        self.export_button.config(state=tk.DISABLED)

        tree_frame = ttk.Frame(press_release_frame)
        tree_frame.pack(fill=tk.BOTH, expand=True)
        scrollbar = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL)
        
        # ✅ [v5.0.1] GUI TreeView는 v5.0의 (Title, Date, Keyword) 3단 컬럼을 유지
        self.press_tree = ttk.Treeview(tree_frame, columns=("Title", "Date", "Keyword"), show="headings", yscrollcommand=scrollbar.set, height=7)
        scrollbar.config(command=self.press_tree.yview)

        self.press_tree.heading("Title", text="보도제목")
        self.press_tree.heading("Date", text="등록일")
        self.press_tree.heading("Keyword", text="검색어 추천")
        
        self.press_tree.column("Title", width=800, anchor=tk.W)
        self.press_tree.column("Date", width=100, anchor=tk.CENTER)
        self.press_tree.column("Keyword", width=120, anchor=tk.CENTER)

        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.press_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # ✅ [v5.0.1] 더블클릭 이벤트는 v5.0의 것을 유지 (검색어 누적 기능)
        self.press_tree.bind("<Double-1>", self.on_press_release_double_click)


        # --- 3. [Step 2] 통합 검색 프레임 ---
        # (이하 v5.0 원본과 동일)
        scraper_frame = ttk.LabelFrame(main_frame, text=" [Step 2] 통합 뉴스 기사 검색 ", padding="10")
        scraper_frame.pack(fill=tk.X, expand=False, pady=5) 

        # --- [v5.0] 검색어 및 N일 (기존과 동일) ---
        ttk.Label(scraper_frame, text="검색어 (쉼표로 구분):").grid(row=0, column=0, padx=5, pady=5, sticky=tk.W)
        
        self.keyword_text_var = tk.StringVar()
        self.keywords_entry = ttk.Entry(scraper_frame, width=60, textvariable=self.keyword_text_var)
        self.keywords_entry.grid(row=0, column=1, padx=5, pady=5, sticky=tk.EW)
        self.keyword_text_var.trace_add("write", self.on_keyword_entry_change)
        
        ttk.Label(scraper_frame, text="최근 N일 (숫자):").grid(row=1, column=0, padx=5, pady=5, sticky=tk.W)
        self.days_entry = ttk.Entry(scraper_frame, width=10)
        self.days_entry.grid(row=1, column=1, padx=5, pady=5, sticky=tk.W)
        
        # --- ✅ [v5.0] 검색 대상 체크박스 프레임 (신규) ---
        check_frame = ttk.LabelFrame(scraper_frame, text=" 검색 대상 (중복 선택 가능) ", padding="5")
        check_frame.grid(row=2, column=0, columnspan=2, padx=5, pady=(5, 10), sticky=tk.EW)
        
        self.fetch_main6_var = tk.BooleanVar(value=True)
        self.fetch_other16_var = tk.BooleanVar(value=True)
        self.fetch_google_var = tk.BooleanVar(value=True)
        
        ttk.Checkbutton(check_frame, text="주요 6대 지방지 (6개)", variable=self.fetch_main6_var).pack(side=tk.LEFT, padx=10)
        ttk.Checkbutton(check_frame, text="기타 지방지 (16개)", variable=self.fetch_other16_var).pack(side=tk.LEFT, padx=10)
        ttk.Checkbutton(check_frame, text="Google 뉴스 (RSS)", variable=self.fetch_google_var).pack(side=tk.LEFT, padx=10)

        # --- ✅ [v5.0] 시작/Word 저장 버튼 영역 ---
        button_frame = ttk.Frame(scraper_frame)
        button_frame.grid(row=0, column=2, rowspan=3, padx=10, pady=5, sticky=tk.NS)

        self.start_button = ttk.Button(button_frame, text="통합 검색 및 엑셀 저장", command=self.start_main_scraper_thread)
        self.start_button.pack(fill=tk.BOTH, expand=True, ipadx=5, ipady=8, pady=(0, 5))

        self.word_button = ttk.Button(button_frame, text="Word 문서로 저장",
                                      command=self.start_word_export_thread,
                                      state=tk.DISABLED)
        self.word_button.pack(fill=tk.BOTH, expand=True, ipadx=5, ipady=8)

        self.start_button.bind("<Return>", lambda event: self.start_main_scraper_thread())
        
        scraper_frame.columnconfigure(1, weight=1)

        # --- 상태바 ---
        self.status_var = tk.StringVar()
        self.status_var.set("준비 완료 (v5.0.1). [Step 1]에서 목록을 가져오거나, [Step 2]에 검색어를 바로 입력하세요.")
        
        status_bar = ttk.Label(root, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W, padding="5", style='Status.TLabel')
        status_bar.pack(side=tk.BOTTOM, fill=tk.X)
        # (이어서 __init__ 메서드 끝)

    def open_link(self, event):
        self.status_var.set(f"{self.goegh_page1_url} 주소로 이동합니다.")
        webbrowser.open_new(self.goegh_page1_url)

    def open_result_folder(self, filename):
        # [v5.0] '6대' 기준 (os.startfile) 및 '기타' 기준 (플랫폼 분기) 통합
        try:
            folder_path = os.path.dirname(os.path.abspath(filename))
            if sys.platform == "win32":
                os.startfile(folder_path)
            elif sys.platform == "darwin": # macOS
                subprocess.Popen(["open", folder_path])
            else: # Linux
                subprocess.Popen(["xdg-open", folder_path])
            print(f"✅ 결과 폴더 자동 열기 시도: {folder_path}")
        except Exception as e:
            print(f"‼️ 결과 폴더 자동 열기 실패: {e}")
            messagebox.showinfo("알림", f"결과 엑셀 파일은 현재 폴더에 저장되었습니다.\n{filename}\n수동으로 폴더를 열어 확인해주세요.")

    def open_file_directly(self, filename):
        try:
            abs_path = os.path.abspath(filename)
            if sys.platform == "win32":
                os.startfile(abs_path)
            elif sys.platform == "darwin":
                subprocess.Popen(["open", abs_path])
            else:
                subprocess.Popen(["xdg-open", abs_path])
            print(f"✅ 파일 열기 시도: {abs_path}")
        except Exception as e:
            print(f"‼️ 파일 열기 실패: {e}")
            messagebox.showinfo("알림", f"파일을 자동으로 열 수 없습니다.\n{filename}\n수동으로 열어주세요.")

    def show_completion_window(self, filename, article_count):
        # [v5.0] '6대 지방지' 기준 팝업 (프로그램을 종료하지 않고 폴더만 열기)
        top = Toplevel(self.root)
        top.title("작업 완료")
        top.geometry("350x180") 
        top.transient(self.root) 
        top.grab_set() 
        main_frame = ttk.Frame(top, padding="15")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        basename = os.path.basename(filename)

        # ✅ [v5.0.1] 팝업 텍스트 분기 (통합검색 / 보도자료 / Word)
        if basename.lower().endswith(".docx"):
            title_text = "✅ Word 문서가 생성되었습니다."
            count_text = f"파일명: {basename}\n(총 {article_count}개 기사 기반)"
            open_label = "문서 열기"
            open_command = lambda: [self.open_file_directly(filename), top.destroy()]
        elif "통합_뉴스검색_" in basename:
            title_text = "✅ 스크래핑이 완료되었습니다."
            count_text = f"파일명: {basename}\n(총 {article_count}개 항목)"
            open_label = "확인 (폴더 열기)"
            open_command = lambda: [self.open_result_folder(filename), top.destroy()]
        else:
            title_text = "✅ 엑셀 저장이 완료되었습니다."
            count_text = f"파일명: {basename}\n(총 {article_count}개 항목)"
            open_label = "확인 (폴더 열기)"
            open_command = lambda: [self.open_result_folder(filename), top.destroy()]

        ttk.Label(main_frame, text=title_text, font=('Malgun Gothic', 12, 'bold')).pack(pady=(5, 0))
        ttk.Label(main_frame, text=count_text).pack(pady=(5, 10))

        open_button = ttk.Button(main_frame, text=open_label, command=open_command)
        open_button.pack(pady=5, ipadx=10, ipady=5)
        
        close_button = ttk.Button(main_frame, text="닫기", command=top.destroy)
        close_button.pack(pady=5)
        
        self.root.wait_window(top) 

    # --- [Step 1] 보도자료 스크레이퍼 스레드 관리 (v5.0.1 수정) ---

    def on_keyword_entry_change(self, *args):
        # (v5.0 원본과 동일)
        try:
            if not self.keyword_text_var.get():
                if self.earliest_date_clicked is not None:
                    print("검색어 창이 비워져 N일 계산 기준을 초기화합니다.")
                    self.earliest_date_clicked = None
                    self.days_entry.delete(0, tk.END) 
                    self.status_var.set("검색어가 초기화되었습니다. 기간도 초기화됩니다.")
        except Exception as e:
            print(f"on_keyword_entry_change 오류: {e}")


    def on_press_release_double_click(self, event):
        # (v5.0 원본과 동일 - GUI가 (Title, Date, Keyword) 기준이므로)
        try:
            item_id = self.press_tree.focus() 
            if not item_id: return
            values = self.press_tree.item(item_id, 'values')
            if not values: return
            
            column_id = self.press_tree.identify_column(event.x)
            
            clicked_date_str = values[1].strip()
            single_keyword_to_add = values[2].strip()
            
            keywords_to_process = [] 

            if column_id == '#2':
                self.status_var.set(f"'{clicked_date_str}'와(과) 동일한 날짜의 모든 키워드를 수집합니다...")
                keywords_set = set() 
                all_item_ids = self.press_tree.get_children()
                
                for an_item_id in all_item_ids:
                    item_values = self.press_tree.item(an_item_id, 'values')
                    if not item_values: continue
                    item_date = item_values[1].strip()
                    item_keyword = item_values[2].strip()
                    
                    if item_date == clicked_date_str and item_keyword:
                        keywords_set.add(item_keyword)
                
                keywords_to_process = list(keywords_set)
                
            else:
                if not single_keyword_to_add:
                    self.status_var.set("이 항목에는 추천 검색어가 없습니다.")
                    return
                keywords_to_process = [single_keyword_to_add]

            if not keywords_to_process:
                self.status_var.set("추가할 키워드가 없습니다.")
                return

            current_text = self.keyword_text_var.get().strip()
            current_keywords = [kw.strip() for kw in current_text.split(',') if kw.strip()]
            
            final_keywords_to_add = []
            for kw in keywords_to_process:
                if kw not in current_keywords:
                    final_keywords_to_add.append(kw)

            if not final_keywords_to_add:
                self.status_var.set(f"선택한 키워드가 이미 모두 목록에 있습니다.")
                return

            # [v5.0] 12개 제한 로직 (3개 파일 공통)
            current_count = len(current_keywords)
            space_available = 12 - current_count
            show_limit_popup = False

            if len(final_keywords_to_add) > space_available:
                if space_available <= 0: 
                    keywords_to_actually_add = []
                else: 
                    keywords_to_actually_add = final_keywords_to_add[:space_available]
                show_limit_popup = True
            else:
                keywords_to_actually_add = final_keywords_to_add
                
            if not keywords_to_actually_add: 
                messagebox.showwarning("입력 초과", "검색어가 이미 12개입니다. (최대 12개)") 
                self.status_var.set("검색어 12개 초과. 추가할 수 없습니다.")
                return
            elif show_limit_popup:
                messagebox.showwarning("입력 초과", "검색어는 최대 12개까지 입력할 수 있습니다. (12개까지만 추가됨)")
                
            current_keywords.extend(keywords_to_actually_add)
            new_text = ", ".join(current_keywords)
            self.keyword_text_var.set(new_text)

            # '최근 N일' 자동 계산
            date_str = clicked_date_str 
            n_days_calculated = None
            try:
                clicked_date = datetime.strptime(date_str, "%Y.%m.%d")
                
                if self.earliest_date_clicked is None or clicked_date < self.earliest_date_clicked:
                    self.earliest_date_clicked = clicked_date
                    
                today = datetime.today()
                delta = today.date() - self.earliest_date_clicked.date()
                n_days = delta.days + 1
                
                if n_days <= 0:
                    n_days = 1
                
                self.days_entry.delete(0, tk.END)
                self.days_entry.insert(0, str(n_days))
                n_days_calculated = n_days

            except ValueError:
                print(f"N일 자동계산 오류: 날짜 형식 오류 ({date_str})")
            except Exception as e_date:
                print(f"N일 자동계산 알 수 없는 오류: {e_date}")

            status_message = (f"키워드 {len(keywords_to_actually_add)}개 추가 (총 {len(current_keywords)}개).")
            if n_days_calculated:
                status_message += f" | 기간 {n_days_calculated}일 자동 설정."
            else:
                status_message += " | 기간 계산 실패."
            self.status_var.set(status_message)
        
        except Exception as e:
            self.status_var.set(f"검색어 복사 실패: {e}")
            traceback.print_exc()

    # ✅ [v5.0.1 신규] (1) 전체 목록 가져오기 (v1.4)
    def start_fetch_press_releases_all(self):
        self.set_step1_buttons_state(tk.DISABLED, "목록 불러오는 중...")
        self.press_tree.delete(*self.press_tree.get_children())
        self.tree_data.clear() 
        self.status_var.set("광주하남 보도자료 목록을 스크래핑합니다 (1페이지)")
        # filter_today=False 전달
        threading.Thread(target=self.run_fetch_press_releases_task, args=(False,), daemon=True).start()

    # ✅ [v5.0.1 신규] (2) 오늘 날짜 목록만 가져오기 (v1.4)
    def start_fetch_press_releases_today(self):
        self.set_step1_buttons_state(tk.DISABLED, "오늘 목록 불러오는 중...")
        self.press_tree.delete(*self.press_tree.get_children())
        self.tree_data.clear() 
        self.status_var.set("오늘 날짜 보도자료 목록을 스크래핑합니다...")
        # filter_today=True 전달
        threading.Thread(target=self.run_fetch_press_releases_task, args=(True,), daemon=True).start()

    # ✅ [v5.0.1 신규] Step 1 버튼 상태 관리 헬퍼
    def set_step1_buttons_state(self, state, main_text=None):
        self.fetch_press_button.config(state=state)
        self.fetch_today_button.config(state=state)
        self.export_button.config(state=state if state == tk.DISABLED else tk.NORMAL) # 엑셀 버튼은 데이터가 있을 때만 활성화됨
        
        if state == tk.DISABLED:
            if "오늘" in (main_text or ""):
                self.fetch_today_button.config(text=main_text)
            else:
                self.fetch_press_button.config(text=main_text)
        else:
            # 복구
            self.fetch_press_button.config(text="보도배포 목록 가져오기(1페이지)")
            today_str = datetime.now().strftime('%Y-%m-%d')
            self.fetch_today_button.config(text=f"'{today_str}' 일자 목록 가져오기")


    def run_fetch_press_releases_task(self, filter_today=False):
        # ✅ [v5.0.2 수정] v1.4의 후처리 로직 적용 (시간 덮어쓰기, 필터링)
        try:
            results = scrape_press_releases(self.goegh_base_url) 
            
            # [요청 1] 등록시간을 스크립트 실행 시간(HH:MM)으로 통일
            current_time_str = datetime.now().strftime('%H:%M')
            for item in results:
                item['time'] = current_time_str

            # [요청 4] 오늘 날짜 필터링 (filter_today=True 일 때)
            if filter_today:
                today_date_str = datetime.now().strftime('%Y.%m.%d')
                print(f"[필터] 오늘 날짜({today_date_str}) 게시글만 필터링합니다.")
                results = [item for item in results if item.get('date') == today_date_str]
            
            # ✅ [v5.0.2 수정] 
            # (v1.3) GUI 표시는 지역별 정렬 <-- 이 로직을 삭제!!
            # def sort_key_region(item):
            # ... (삭제) ...
            # results.sort(key=sort_key_region)
            #
            # -> GUI는 크롤링한 원본 순서(최신순) 그대로 표시합니다.
            
            self.root.after(0, self.update_press_release_treeview, results, filter_today)
            
        except Exception as e:
            print(f"‼️ 보도자료 스크래핑 스레드 오류: {e}")
            traceback.print_exc()
            self.root.after(0, self.update_press_release_treeview, [], filter_today)

    def update_press_release_treeview(self, results, filter_today=False):
        # (이 함수는 v5.0.1과 동일 - 수정 없음)
        # ✅ [v5.0.1 수정] 
        # GUI는 (Title, Date, Keyword) 유지, 엑셀용 tree_data는 원본(results) 저장
        self.press_tree.delete(*self.press_tree.get_children())
        self.tree_data.clear()
        
        if not results:
            status_msg = "오늘 날짜 보도자료가 없습니다." if filter_today else "보도자료 목록을 가져오지 못했거나, 목록이 비어있습니다."
            self.status_var.set(status_msg)
        else:
            for item in results:
                # [v5.0.1] GUI에는 v5.0 형식 (제목, 날짜, 추천키워드) 사용
                values_gui = (item["title"], item["date"], item["keyword_gui"])
                
                item_id = self.press_tree.insert("", tk.END, values=values_gui)
                # [v5.0.1] 엑셀용 데이터는 원본 item을 저장 (v1.4)
                self.tree_data[item_id] = item 
            
            # ✅ [v5.0.2 수정] 상태바 메시지 변경 (정렬 문구 삭제)
            status_msg = f"오늘 날짜 보도자료 {len(results)}개 로드 완료." if filter_today else f"보도자료 {len(results)}개 로드 완료. (최신순)"
            self.status_var.set(status_msg)
            self.export_button.config(state=tk.NORMAL) # 엑셀 버튼 활성화

        # [v5.0.1] Step 1 버튼 상태 복구
        self.set_step1_buttons_state(tk.NORMAL)
        
        # [v5.0.1] Step 2 버튼은 항상 활성화 (Step 1과 독립적이므로)
        self.start_button.config(state=tk.NORMAL)

    # --- ✅ [v5.0.1 신규] Step 1 엑셀 저장 (v1.4 뷰어 기준) ---
    def export_to_excel(self):
        if not self.tree_data:
            messagebox.showwarning("데이터 없음", "먼저 '목록 가져오기'를 실행해주세요.")
            return

        today_str = datetime.today().strftime('%Y%m%d_%H%M')
        filename = filedialog.asksaveasfilename(
            initialfile=f"광주하남_보도자료_{today_str}.xlsx",
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")]
        )

        if not filename:
            self.status_var.set("엑셀 저장이 취소되었습니다.")
            return

        self.status_var.set(f"엑셀 파일 저장 중... ({os.path.basename(filename)})")
        # 모든 버튼 비활성화 (저장 중)
        self.set_step1_buttons_state(tk.DISABLED, "엑셀 저장 중...")
        self.start_button.config(state=tk.DISABLED)

        try:
            # tree_data는 item_id: item 딕셔너리. 값만 리스트로 변환
            # (v5.0.2) data_list는 GUI와 동일한 '크롤링 순서(최신순)'를 유지
            data_list = list(self.tree_data.values())
            
            # ✅ [v5.0.2 수정]
            # (v1.3) 엑셀은 날짜순 정렬 (GUI 정렬과 다름) <-- 이 로직을 삭제
            # data_list.sort(key=lambda item: (item.get('date', ''), item.get('region', '')))
            
            # ✅ [v5.0.2 신규] 요청사항: 엑셀은 (교육지원청 > 광주 > 하남) 순서로 정렬.
            # (Python의 sort()는 stable하므로, 기존의 '크롤링 순서(최신순)'가 2차 정렬로 유지됨)
            def sort_key_excel_region(item):
                region = item.get('region', '교육지원청') # 기본값을 '교육지원청'으로
                if region == '교육지원청':
                    return 0
                if region == '광주':
                    return 1
                if region == '하남':
                    return 2
                return 3 # 기타 (혹시 모를 경우)

            data_list.sort(key=sort_key_excel_region)
            
            processed_list = []
            for i, item in enumerate(data_list):
                recommend_val = "★" if i == 0 else item.get("priority", "")
                
                processed_list.append({
                    "recommend": recommend_val,
                    "yeonbeon": i + 1,
                    "region": item.get("region", ""),
                    "date": item.get("date", ""),
                    "time": item.get("time", ""), # 스크래핑 시점의 시간
                    "title": item.get("title", ""),
                    "institution": item.get("institution", ""),
                    "notes": item.get("notes", "")
                })
            
            df = pd.DataFrame(processed_list)

            column_keys = ["recommend", "yeonbeon", "region", "date", "time", "title", "institution", "notes"]
            header_names = ["추천", "연번", "지역", "등록일", "등록시간", "보도제목", "등록기관", "비고"]
            df = df[column_keys]

            with pd.ExcelWriter(filename, engine="xlsxwriter") as writer:
                workbook = writer.book
                worksheet = workbook.add_worksheet("보도자료") 

                # (v1.3) 서식 정의
                header_format = workbook.add_format({'bold': True, 'text_wrap': True, 'valign': 'vcenter', 'align': 'center', 'fg_color': '#D7E4BC', 'border': 1})
                title_format = workbook.add_format({'border': 1, 'align': 'left', 'valign': 'vcenter', 'text_wrap': True})
                center_format = workbook.add_format({'border': 1, 'align': 'center', 'valign': 'vcenter'})
                left_format = workbook.add_format({'border': 1, 'align': 'left', 'valign': 'vcenter'})

                for col_num, header in enumerate(header_names):
                    worksheet.write(0, col_num, header, header_format)
                
                worksheet.set_default_row(30) 
                worksheet.set_row(0, 20) 

                for row_num, row in enumerate(df.itertuples(index=False), start=1):
                    worksheet.write(row_num, 0, row.recommend, center_format) # A: 추천
                    worksheet.write(row_num, 1, row.yeonbeon, center_format)  # B: 연번
                    worksheet.write(row_num, 2, row.region, center_format)    # C: 지역
                    worksheet.write(row_num, 3, row.date, center_format)      # D: 등록일
                    worksheet.write(row_num, 4, row.time, center_format)      # E: 등록시간
                    worksheet.write(row_num, 5, row.title, title_format)      # F: 보도제목
                    worksheet.write(row_num, 6, row.institution, left_format) # G: 등록기관
                    worksheet.write(row_num, 7, row.notes, left_format)       # H: 비고

                # --- 5-5. [v1.4] 컬럼 너비 설정 ---
                worksheet.set_column("A:A", 5)  # 추천
                worksheet.set_column("B:B", 5)  # 연번
                worksheet.set_column("C:C", 6)  # 지역
                worksheet.set_column("D:D", 12) # 등록일
                worksheet.set_column("E:E", 8) # 등록시간
                worksheet.set_column("F:F", 70) # 보도제목
                worksheet.set_column("G:G", 18) # 등록기관
                worksheet.set_column("H:H", 5) # 비고
                
                worksheet.freeze_panes(1, 0) 
                worksheet.autofilter(0, 0, len(df), len(header_names) - 1)

            print(f"\n✅ 저장 완료: {filename}")
            
            self.root.after(0, lambda: [
                self.show_completion_window(filename, len(df)),
                self.status_var.set(f"저장 완료: {os.path.basename(filename)}. 새 작업 가능."),
                # (v1.4) [수정] 엑셀 저장 완료 후 모든 버튼 활성화
                self.set_step1_buttons_state(tk.NORMAL),
                self.start_button.config(state=tk.NORMAL)
            ])

        except Exception as e:
            print(f"\n❌ 엑셀 저장 오류 발생:")
            traceback.print_exc()
            self.root.after(0, lambda e=e: [ 
                messagebox.showerror("엑셀 저장 오류", f"파일 저장 중 오류가 발생했습니다:\n{e}\n\n(파일이 열려있는지 확인하세요)"),
                self.status_var.set(f"엑셀 저장 오류: {e}."),
                # (v1.4) [수정] 엑셀 저장 오류 시에도 모든 버튼 활성화
                self.set_step1_buttons_state(tk.NORMAL),
                self.start_button.config(state=tk.NORMAL)
            ])

    # --- [Step 2] 6대 지방지 스크레이퍼 스레드 관리 (✅ v5.0 원본 유지) ---
    
    def start_main_scraper_thread(self):
        self.status_var.set("입력 값 검증 중...")
        
        try:
            days_limit = int(self.days_entry.get())
            if days_limit <= 0: raise ValueError
        except ValueError:
            self.status_var.set("입력 오류: '최근 N일'은 1 이상의 숫자로 입력해야 합니다.")
            messagebox.showerror("입력 오류", "최근 N일은 1 이상의 숫자로 입력해야 합니다.")
            return

        keywords = [kw.strip() for kw in self.keyword_text_var.get().split(",") if kw.strip()]
        if not keywords:
            self.status_var.set("입력 오류: '검색어'를 1개 이상 입력해야 합니다.")
            messagebox.showerror("입력 오류", "검색어를 1개 이상 입력해야 합니다.")
            return
        
        if len(keywords) > 12:
            messagebox.showwarning("입력 확인", "검색어가 12개를 초과했습니다. (최대 12개)")
            keywords = keywords[:12]
            self.keyword_text_var.set(", ".join(keywords))
            
        # ✅ [v5.0] 체크박스 값 읽기
        fetch_main6 = self.fetch_main6_var.get()
        fetch_other16 = self.fetch_other16_var.get()
        fetch_google = self.fetch_google_var.get()
        
        if not fetch_main6 and not fetch_other16 and not fetch_google:
            self.status_var.set("입력 오류: '검색 대상'을 1개 이상 선택해야 합니다.")
            messagebox.showerror("입력 오류", "검색 대상을 1개 이상 선택해야 합니다.")
            return
        
        self.status_var.set(f"[{keywords[0]} 외 {len(keywords)-1}개] (최근 {days_limit}일) 통합 검색을 시작합니다...")
        
        # ✅ [v5.0.1 수정] Step 2 시작 시 모든 버튼 비활성화
        self.start_button.config(state=tk.DISABLED)
        self.word_button.config(state=tk.DISABLED)
        self.latest_articles_df = None
        self.set_step1_buttons_state(tk.DISABLED, "Step 2 검색 실행 중...")
        
        threading.Thread(target=self.run_main_scraper_task, 
                         args=(keywords, days_limit, fetch_main6, fetch_other16, fetch_google), 
                         daemon=True).start()

    def run_main_scraper_task(self, keywords, days_limit, fetch_main6, fetch_other16, fetch_google):
        # (v5.0 원본과 동일)
        filename = ""
        all_articles = []
        
        # (name, function, type)
        functions_map_main6 = [
            ("기호일보", fetch_kiho_multi, "group"),
            ("경기일보", fetch_kyeonggi_multi, "group"),
            ("경인일보", fetch_kyeongin_multi, "group"),
            ("경기신문", fetch_kgnews_multi, "group"),
            ("인천일보", fetch_incheonilbo_multi, "group"),
            ("중부일보", fetch_joongbu_multi, "group"),
        ]
        
        functions_map_other16 = [
            ("경기핫타임스", fetch_kghottimes_multi, "group"),
            ("하남타임즈", fetch_hanamtimes_multi, "group"),
            ("시티뉴스", fetch_ctnews_multi, "group"),
            ("굿타임즈", fetch_goodtimes_multi, "group"),
            ("하나로신문", fetch_hnrsm_multi, "group"),
            ("하남신문", fetch_ehanam_multi, "group"),
            ("광주신문", fetch_gjilbo_multi, "group"),
            ("중부시사신문", fetch_jungbusisa_multi, "group"),
            ("포탈뉴스통신", fetch_portalnews_multi, "group"),
            ("경기연합뉴스", fetch_kgyonhap_multi, "group"),
            ("교차로저널", fetch_kocus_multi, "group"),
            ("비전21뉴스", fetch_vision21_multi, "group"),
            ("뉴스투데이24", fetch_newstoday24_multi, "group"),
            ("투데이광주하남", fetch_tgh_multi, "group"),
            ("미디어리포트", fetch_mediareport_multi, "group"),
            ("수도권일보", fetch_sudokwon_multi, "group"),
        ]
        
        functions_map_google = [
            ("Google뉴스(RSS)", fetch_google_news_feed, "individual"),
        ]
        
        try:
            self.root.after(0, lambda: self.status_var.set("작업 스레드 시작... 검색 기간 설정 중..."))
            
            today = datetime.today()
            date_limit_dt = today - timedelta(days=days_limit)
            date_limit = date_limit_dt.replace(hour=0, minute=0, second=0, microsecond=0)
            
            print(f"--- 검색 시작 (키워드 {len(keywords)}개, 검색 기간: {date_limit.strftime('%Y-%m-%d')} ~ {today.strftime('%Y-%m-%d')}) ---")
            
            # ✅ [v5.0] 통합 엑셀 파일명
            filename = f"통합_뉴스검색_{today.strftime('%Y%m%d_%H%M')}.xlsx"
            
            self.root.after(0, lambda: self.status_var.set(f"엑셀 파일명 생성: {filename}"))

            futures_to_run = {}
            
            max_workers_selenium = 8
            max_workers_google = 4
            
            with ThreadPoolExecutor(max_workers=max_workers_selenium + max_workers_google) as executor:
                
                # --- 1. 주요 6대 지방지 작업 제출 (group) ---
                if fetch_main6:
                    print("--- [v5.0] 주요 6대 지방지 검색 시작 ---")
                    for name, func, f_type in functions_map_main6:
                        future = executor.submit(func, keywords, date_limit, days_limit) 
                        futures_to_run[future] = (name, "주요6대")

                # --- 2. 기타 16대 지방지 작업 제출 (group) ---
                if fetch_other16:
                    print("--- [v5.0] 기타 16개 지방지 검색 시작 ---")
                    for name, func, f_type in functions_map_other16:
                        future = executor.submit(func, keywords, date_limit, days_limit) 
                        futures_to_run[future] = (name, "기타16대")
                
                # --- 3. Google 뉴스 작업 제출 (individual) ---
                if fetch_google:
                    print("--- [v5.0] Google 뉴스 검색 시작 ---")
                    for name, func, f_type in functions_map_google:
                        # ✅ [v5.0.3 수정] individual 타입은 func이 키워드 목록을 직접 처리하도록 한번만 제출
                        if f_type == "group" or f_type == "individual":
                            future = executor.submit(func, keywords, date_limit, days_limit)
                            futures_to_run[future] = (name, "Google")
                
                # --- 4. 통합 결과 처리 ---
                total_futures = len(futures_to_run)
                completed_count = 0
                
                print(f"--- [v5.0] 총 {total_futures}개의 작업을 병렬 처리합니다. ---")
                
                for future in as_completed(futures_to_run):
                    completed_count += 1
                    source_name, source_group = futures_to_run[future] 
                    
                    status_prefix = f"[ ({completed_count}/{total_futures}) {source_group} ]"
                    
                    try:
                        items = future.result() # items는 list
                        
                        if not items:
                            print(f"📪 [{source_name}] 작업 반환값이 비어있음.")
                            self.root.after(0, lambda p=status_prefix, s=source_name: 
                                            self.status_var.set(f"{p} 📪 [{s}] 결과 없음."))
                            continue

                        article_count = sum(1 for item in items if "없음" not in item.get("보도일", "") and "오류" not in item.get("보도일", ""))
                        
                        if article_count == 0:
                            self.root.after(0, lambda p=status_prefix, s=source_name: 
                                            self.status_var.set(f"{p} 📪 [{s}] 기사 없음."))
                        else:
                            self.root.after(0, lambda p=status_prefix, s=source_name, c=article_count: 
                                            self.status_var.set(f"{p} 📰 [{s}] {c}건 발견!"))
                        
                        all_articles.extend(items)

                    except Exception as e:
                        print(f"--- ‼️ {source_name} 처리 중 심각한 오류 발생 ---")
                        traceback.print_exc() 
                        self.root.after(0, lambda p=status_prefix, s=source_name: 
                                        self.status_var.set(f"{p} ❌ [{s}] 오류 발생."))
                
            # (as_completed 루프가 끝난 후)
            
            if not all_articles:
                print("\n모든 검색어에 대해 수집된 기사가 없습니다.")
                self.root.after(0, lambda: [
                    messagebox.showinfo("검색 완료", "모든 검색어에 대해 수집된 기사가 없습니다."),
                    self.status_var.set("검색 완료. 수집된 기사가 없습니다. (새 검색 가능)"),
                    # ✅ [v5.0.1] 모든 버튼 활성화
                    self.start_button.config(state=tk.NORMAL),
                    self.set_step1_buttons_state(tk.NORMAL),
                    self.word_button.config(state=tk.DISABLED)
                ])
                return

            self.root.after(0, lambda: self.status_var.set(f"총 {len(all_articles)}개 항목 수집 완료. 데이터 정렬 중..."))
            
            df = pd.DataFrame(all_articles)
            
            # '없음'/'오류' 항목은 정렬 시 하단으로 보내기 위해 임시 날짜(1900년) 부여
            today_str = datetime.today().strftime('%Y-%m-%d')
            
            def clean_date_for_sorting(date_str):
                if isinstance(date_str, str):
                    if "없음" in date_str or "오류" in date_str or "확인 바람" in date_str:
                        return pd.Timestamp('1900-01-01')
                    if re.match(r"^\d{4}-\d{2}-\d{2}$", date_str):
                        return pd.to_datetime(date_str, format='%Y-%m-%d', errors='coerce')
                return pd.NaT # 그 외 파싱 불가 날짜

            df['보도일_dt'] = df['보도일'].apply(clean_date_for_sorting)
            
            # ✅ [v5.0] 요청 사항: '보도제목' (오름차순)으로 정렬. 
            df = df.sort_values(by=['보도제목', '보도일_dt', '보도매체'], ascending=[True, False, True])
            
            df = df.drop(columns=['보도일_dt']) 
            
            # [v5.0] 엑셀 열 순서 (공통)
            df = df[["보도일", "보도매체", "보도제목", "링크", "검색어"]]
            self.latest_articles_df = df.copy()

            print("\n💾 엑셀 파일 저장 중...")
            self.root.after(0, lambda: self.status_var.set(f"데이터 정렬 완료. 엑셀 파일({filename}) 저장 중..."))
            
            with pd.ExcelWriter(filename, engine="xlsxwriter") as writer:
                df.to_excel(writer, sheet_name="통합결과", index=False, startrow=1, header=False)
                workbook = writer.book
                worksheet = writer.sheets["통합결과"]
                
                header_format = workbook.add_format({
                    'bold': True, 'text_wrap': True, 'valign': 'vcenter',
                    'fg_color': '#D7E4BC', 'border': 1 
                })
                
                headers = ["보도일", "보도매체", "보도제목", "링크", "검색어"]
                for col_num, header in enumerate(headers):
                    worksheet.write(0, col_num, header, header_format)
                
                worksheet.freeze_panes(1, 0)
                url_format = workbook.add_format({'font_color': 'blue', 'underline': 1})
                
                for row_num, row in enumerate(df.itertuples(index=False), start=1):
                    worksheet.write(row_num, 0, row.보도일)
                    worksheet.write(row_num, 1, row.보도매체)
                    
                    if (row.링크 and isinstance(row.링크, str) and 
                        row.링크.startswith("http") and 
                        "없음" not in row.보도일 and 
                        "오류" not in row.보도일 and
                        "확인 바람" not in row.보도일):
                        try:
                            worksheet.write_url(row_num, 2, row.링크, string=row.보도제목, cell_format=url_format)
                        except ValueError as excel_url_error:
                            print(f"엑셀 URL 쓰기 오류 (텍스트로 대체): {excel_url_error} | 링크: {row.링크}")
                            worksheet.write(row_num, 2, row.보도제목) # 오류 시 제목만 텍스트로
                    else:
                        worksheet.write(row_num, 2, row.보도제목) 
                        
                    worksheet.write(row_num, 3, row.링크) 
                    worksheet.write(row_num, 4, row.검색어)

                worksheet.set_column("A:A", 12) # 보도일
                worksheet.set_column("B:B", 18) # 보도매체
                worksheet.set_column("C:C", 70) # 보도제목
                worksheet.set_column("D:D", 40) # 링크
                worksheet.set_column("E:E", 15) # 검색어
                worksheet.autofilter(0, 0, len(df), len(headers) - 1)

            print(f"\n✅ 저장 완료: {filename}")
            
            # ✅ [v5.0.1] 완료 팝업 및 모든 버튼 활성화
            self.root.after(0, lambda: [
                self.show_completion_window(filename, len(all_articles)),
                self.status_var.set(f"저장 완료: {filename}. (총 {len(all_articles)}개). 새 검색 가능."),
                self.start_button.config(state=tk.NORMAL),
                self.set_step1_buttons_state(tk.NORMAL),
                self.word_button.config(state=tk.NORMAL)
            ])

        except Exception as e:
            print("\n❌ 예외 발생:")
            traceback.print_exc()
            self.root.after(0, lambda e=e: [
                messagebox.showerror("치명적 오류", f"작업 중 오류가 발생했습니다:\n{e}\n\n(자세한 내용은 콘솔 창을 확인하세요)"),
                self.status_var.set(f"오류 발생: {e}. (콘솔 창 확인)"),
                # ✅ [v5.0.1] 모든 버튼 활성화
                self.start_button.config(state=tk.NORMAL),
                self.set_step1_buttons_state(tk.NORMAL),
                self.word_button.config(state=tk.DISABLED)
            ])

    # --- Word 문서 생성 기능 ---

    def start_word_export_thread(self):
        if not isinstance(self.latest_articles_df, pd.DataFrame) or self.latest_articles_df.empty:
            messagebox.showinfo("Word 문서 생성", "먼저 '통합 검색 및 엑셀 저장'을 실행하여 데이터를 수집해 주세요.")
            return

        df_snapshot = self.latest_articles_df.copy()
        self.word_button.config(state=tk.DISABLED)
        self.status_var.set("Word 문서를 생성 준비 중입니다...")

        threading.Thread(target=self.run_word_export_task, args=(df_snapshot,), daemon=True).start()

    def run_word_export_task(self, df_snapshot):
        driver = None
        try:
            self.root.after(0, lambda: self.status_var.set("Word 문서 생성 중... (브라우저 드라이버 초기화)"))
            driver = setup_driver_compatible(headless=True)
            if not driver:
                raise Exception("웹 드라이버를 시작할 수 없습니다.")

            self.root.after(0, lambda: self.status_var.set("Word 문서 생성 중... (링크 변환 및 문서 생성)"))
            info = self._create_word_summary_document(df_snapshot, driver)
            
            self.root.after(0, lambda info=info: [
                self.show_completion_window(info["filename"], info["article_count"]),
                self.status_var.set(f"Word 문서 저장 완료: {os.path.basename(info['filename'])}"),
                self.word_button.config(state=tk.NORMAL)
            ])
        except Exception as e:
            print("\n❌ Word 문서 생성 오류:")
            traceback.print_exc()
            self.root.after(0, lambda e=e: [
                messagebox.showerror("Word 저장 오류", f"Word 문서 생성 중 오류가 발생했습니다:\n{e}"),
                self.status_var.set(f"Word 저장 오류: {e}"),
                self.word_button.config(state=tk.NORMAL)
            ])
        finally:
            if driver:
                print("Word 문서 생성을 위한 임시 드라이버를 종료합니다.")
                driver.quit()

    def _create_word_summary_document(self, df, driver):
        latest_dt, latest_label = self._get_latest_date_info(df)
        doc_filename = f"일일보도_요약_{latest_dt.strftime('%Y%m%d_%H%M')}.docx"
        output_path = os.path.abspath(doc_filename)

        document = Document()
        self._configure_word_styles(document)

        title_paragraph = document.add_heading(f"📰 광주하남교육 뉴스 보도기사 ({latest_label} 자)", level=1)
        self._set_paragraph_font(title_paragraph, 18)

        for _ in range(2):
            spacer = document.add_paragraph("")
            self._set_paragraph_font(spacer, 13)

        grouped_entries = self._group_articles_for_word(df)

        for entry in grouped_entries:
            line_paragraph = document.add_paragraph(f"- {entry['title']}  {entry['media_text']}")
            self._set_paragraph_font(line_paragraph, 13)
            line_paragraph.paragraph_format.space_after = Pt(4)

            if entry["link"]:
                clean_link = self._resolve_google_news_url(entry["link"], driver)
                link_paragraph = document.add_paragraph("   ↪ ")
                self._set_paragraph_font(link_paragraph, 11)
                self._add_hyperlink(link_paragraph, clean_link, clean_link)
                link_paragraph.paragraph_format.space_after = Pt(6)

            document.add_paragraph("")

        document.save(output_path)
        return {"filename": output_path, "article_count": len(df)}

    def _configure_word_styles(self, document):
        normal_style = document.styles["Normal"]
        normal_style.font.name = "Malgun Gothic"
        normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        normal_style.font.size = Pt(12)

        heading_style = document.styles["Heading 1"]
        heading_style.font.name = "Malgun Gothic"
        heading_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")

    def _set_paragraph_font(self, paragraph, size):
        runs = paragraph.runs
        if not runs:
            runs = [paragraph.add_run("")]
        for run in runs:
            run.font.name = "Malgun Gothic"
            try:
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
            except AttributeError:
                pass
            run.font.size = Pt(size)

    def _group_articles_for_word(self, df):
        grouped = []
        for _, row in df.iterrows():
            title = str(row.get("보도제목", "") or "").strip()
            if not title:
                continue
            media = str(row.get("보도매체", "") or "").strip()
            link = str(row.get("링크", "") or "").strip()

            matched_group = None
            for group in grouped:
                existing_title = group["titles"][0]
                if self._calculate_title_similarity(title, existing_title) >= 0.6:
                    matched_group = group
                    break

            if not matched_group:
                matched_group = {"titles": [title], "medias": [], "links": []}
                grouped.append(matched_group)
            elif title not in matched_group["titles"]:
                matched_group["titles"].append(title)

            matched_group["medias"].append(media)
            matched_group["links"].append(link)

        structured = []
        for group in grouped:
            if not group["titles"]:
                continue
            representative_media, media_text = self._determine_media_text(group["medias"])
            final_link = self._pick_representative_link(group["medias"], group["links"], representative_media)
            structured.append({
                "title": group["titles"][0],
                "media_text": media_text,
                "link": final_link
            })

        return structured

    def _determine_media_text(self, medias):
        unique_media = []
        for media in medias:
            media = media.strip() if isinstance(media, str) else ""
            if media and media not in unique_media:
                unique_media.append(media)

        non_hanaro = [m for m in unique_media if m != "하나로신문"]

        if not unique_media:
            return "", "<매체 정보 없음>"

        if len(unique_media) == 1 and unique_media[0] == "하나로신문":
            return "하나로신문", "<하나로신문>"

        if not non_hanaro:
            return "하나로신문", "<하나로신문>"

        if len(non_hanaro) == 1:
            return non_hanaro[0], f"<{non_hanaro[0]}>"

        return non_hanaro[0], f"<{non_hanaro[0]} 외 {len(non_hanaro) - 1}건>"

    def _pick_representative_link(self, medias, links, representative_media):
        if representative_media:
            for media, link in zip(medias, links):
                if media == representative_media and link:
                    return link
        for link in links:
            if link:
                return link
        return ""

    def _get_latest_date_info(self, df):
        date_series = pd.to_datetime(df["보도일"], errors="coerce")
        latest_dt = date_series.dropna().max()
        if pd.isna(latest_dt):
            latest_dt = datetime.today()
        if isinstance(latest_dt, pd.Timestamp):
            latest_dt = latest_dt.to_pydatetime()
        label = f"{latest_dt.year}. {latest_dt.month}. {latest_dt.day}."
        return latest_dt, label

    def _resolve_google_news_url(self, url, driver):
        if not url or not driver:
            return url or ""
            
        url = url.strip()
        # Check if it's a google news url that needs resolving
        if "news.google.com/rss/articles/" not in url:
            return url

        try:
            print(f"Google 뉴스 링크 감지, Selenium으로 최종 주소 확인: {url}")
            driver.get(url)
            # Wait up to 10 seconds for the URL to change from the original google one.
            # This handles JS redirects.
            WebDriverWait(driver, 10).until(
                lambda d: d.current_url != url and "news.google.com" not in d.current_url
            )
            final_url = driver.current_url
            print(f"링크 변환 성공 (Selenium): {url} -> {final_url}")
            return final_url
        except TimeoutException:
            # If the URL doesn't change, it might be a direct link or a failed redirect.
            # Return the URL we have, which might still be the google one.
            current_url = driver.current_url
            print(f"Selenium 타임아웃: URL이 변경되지 않았습니다. 현재 URL 반환: {current_url}")
            return current_url
        except Exception as e:
            print(f"Selenium으로 링크 변환 중 오류 발생 (원본 URL 반환): {e}")
            return url # Fallback to original URL on error

    def _calculate_title_similarity(self, s1, s2):
        if not s1 or not s2:
            return 0.0
        distance = self._levenshtein_distance(s1, s2)
        max_len = max(len(s1), len(s2))
        return 1.0 - distance / max_len if max_len else 1.0

    def _levenshtein_distance(self, a, b):
        if a == b:
            return 0
        if not a:
            return len(b)
        if not b:
            return len(a)

        previous_row = list(range(len(b) + 1))
        for i, ca in enumerate(a, start=1):
            current_row = [i]
            for j, cb in enumerate(b, start=1):
                insertions = previous_row[j] + 1
                deletions = current_row[j - 1] + 1
                substitutions = previous_row[j - 1] + (ca != cb)
                current_row.append(min(insertions, deletions, substitutions))
            previous_row = current_row
        return previous_row[-1]

    def _add_hyperlink(self, paragraph, url, text):
        if not url:
            paragraph.add_run(text)
            return

        part = paragraph.part
        r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)
        new_run.append(rPr)

        text_elem = OxmlElement("w:t")
        text_elem.text = text
        new_run.append(text_elem)
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)

# =============================================================================
# 8. GUI 시작 (v5.0.1) -> (v5.0.2로 이름만 변경)
# =============================================================================
if __name__ == "__main__":
    try:
        # ✅ [v5.0.2 수정] 버전명 변경
        print("프로그램 시작... (v5.0.2) GUI를 시작합니다.") 
        print("(첫 실행 시 WebDriverManager가 크롬 드라이버를 다운로드할 수 있습니다...)")
        
        get_driver_path() 
        print("웹 드라이버 준비 완료. GUI를 시작합니다.")
        
        root = tk.Tk()
        app = NewsScraperApp(root)
        root.mainloop()
        
    except Exception as e:
        print("\n❌ GUI 시작 오류 발생:")
        traceback.print_exc()
        try:
            temp_root = tk.Tk()
            temp_root.withdraw()
            messagebox.showerror("GUI 시작 오류", 
                               f"프로그램 시작 중 치명적인 오류가 발생했습니다.\n\n"
                               f"오류: {e}\n"
                               f"(자세한 내용은 콘솔 창을 확인해주세요.)")
            temp_root.destroy()
        except tk.TclError:
             print("--- GUI(Tkinter) 라이브러리 로드 실패 ---")
             input("엔터 키를 누르면 종료됩니다...")
