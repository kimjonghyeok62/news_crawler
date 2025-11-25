from flask import Flask, render_template, request, jsonify, send_file
import os
import pandas as pd
from datetime import datetime
import threading
from crawler_service import scrape_press_releases, run_crawler_task, job_manager
from docx import Document
from docx.shared import Pt
import io

app = Flask(__name__)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/crawl/press', methods=['POST'])
def crawl_press():
    # Step 1: Press Releases
    try:
        base_url = "https://www.goegh.kr/goegh/na/ntt/selectNttList.do?mi=8686&bbsId=5041"
        results = scrape_press_releases(base_url)
        
        # Filter for today if requested
        filter_today = request.json.get('filter_today', False)
        if filter_today:
            today_str = datetime.now().strftime('%Y.%m.%d')
            results = [r for r in results if r.get('date') == today_str]
            
        return jsonify({"status": "success", "data": results})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/api/crawl/news', methods=['POST'])
def crawl_news():
    # Step 2: News Crawling (Async)
    data = request.json
    keywords = data.get('keywords', [])
    days_limit = int(data.get('days_limit', 1))
    sources = data.get('sources', ['google']) # ['google', 'major6', 'other16']
    
    if not keywords:
        return jsonify({"status": "error", "message": "No keywords provided"}), 400
        
    job_id = job_manager.create_job()
    
    # Run in background thread
    thread = threading.Thread(target=run_crawler_task, args=(job_id, keywords, days_limit, sources))
    thread.start()
    
    return jsonify({"status": "success", "job_id": job_id})

@app.route('/api/status/<job_id>', methods=['GET'])
def get_status(job_id):
    job = job_manager.get_job(job_id)
    if not job:
        return jsonify({"status": "error", "message": "Job not found"}), 404
    return jsonify(job)

@app.route('/api/export/excel', methods=['POST'])
def export_excel():
    data = request.json.get('data', [])
    filename = f"export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    
    df = pd.DataFrame(data)
    
    # Create Excel in memory
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Sheet1')
    output.seek(0)
    
    return send_file(output, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

@app.route('/api/export/word', methods=['POST'])
def export_word():
    data = request.json.get('data', [])
    filename = f"summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    doc = Document()
    doc.add_heading('News Summary', 0)
    
    for item in data:
        p = doc.add_paragraph()
        p.add_run(f"[{item.get('보도매체', 'Unknown')}] ").bold = True
        p.add_run(item.get('보도제목', 'No Title'))
        p.add_run(f"\nLink: {item.get('링크', '')}").italic = True
        p.add_run(f"\nDate: {item.get('보도일', '')}")
        
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return send_file(output, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
